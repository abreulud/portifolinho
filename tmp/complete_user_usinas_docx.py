from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


SOURCE = Path(r"C:\Users\Ludmilla\Downloads\Projeto Usinas de Geração de Energia.docx")
OUTPUT = Path(r"C:\Users\Ludmilla\Documents\WebsitePortifolio\outputs\Projeto Usinas de Geração de Energia - Complementado.docx")
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document(SOURCE)


def find_paragraph(text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text.strip():
            return paragraph
    raise ValueError(f"Parágrafo não encontrado: {text}")


def find_paragraph_with_style(text, style_name):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text.strip() and paragraph.style.name == style_name:
            return paragraph
    raise ValueError(f"Parágrafo não encontrado: {text} ({style_name})")


def replace_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_after(paragraph, text="", style="normal"):
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    if style:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


bullet_template = find_paragraph("Requisitos")


def insert_bullet_after(paragraph, text):
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    if bullet_template._p.pPr is not None:
        inserted._p.insert(0, deepcopy(bullet_template._p.pPr))
    inserted.add_run(text)
    return inserted


def insert_label_after(paragraph, label, text=""):
    inserted = insert_after(paragraph)
    inserted.paragraph_format.keep_with_next = not bool(text)
    run = inserted.add_run(label)
    run.bold = True
    if text:
        inserted.add_run(text)
    return inserted


def add_activity_content(heading_text, description, activities, estimate):
    anchor = find_paragraph(heading_text)
    anchor = insert_after(anchor, description)
    anchor = insert_label_after(anchor, "Atividades:")
    for activity in activities:
        anchor = insert_bullet_after(anchor, activity)
    anchor = insert_label_after(anchor, "Estimativa: ", estimate)
    return anchor


# Ajustes textuais pontuais
replace_text(
    find_paragraph("Este plano tem como objetivo definir as atividades para desenvolver, testar e transferir uma solução tecnológica de segurança criptográfica em usinas de geração de energia."),
    "Este plano tem como objetivo definir as atividades necessárias para conceber, desenvolver, testar e transferir soluções de segurança criptográfica aplicáveis a usinas de geração de energia.",
)
replace_text(
    find_paragraph("O projeto atual envolve duas frentes tecnológicas: A primeira trata do detalhamento da arquitetura, inventário, cryptoagility e proxy. A segunda trata da avaliação e integração com Distribuição Quânticas de Chaves (QKD)."),
    "O projeto está organizado em duas frentes tecnológicas. A primeira trata de arquitetura, inventário, criptoagilidade e proxy. A segunda trata da avaliação e integração com Distribuição Quântica de Chaves (QKD).",
)
replace_text(find_paragraph("1. Benchmark e Informacional"), "1. Benchmarking e Informacional")
replace_text(find_paragraph("Requisitos"), "Levantamento de requisitos")
replace_text(find_paragraph("Arquiteturas"), "Levantamento de arquiteturas")
replace_text(find_paragraph("Treinamento e capacitação de equipes"), "Treinamento e capacitação das equipes")
replace_text(find_paragraph("2.2 Governança e Cripto Agility"), "2.2 Governança e Criptoagilidade")
replace_text(find_paragraph("3.3 Desacoplamento das Aplicações Criptograficas"), "3.3 Desacoplamento das Aplicações da Criptografia")
replace_text(find_paragraph("4.1 Testes funcionais, Integração e Interoperabilidade"), "4.1 Testes Funcionais, de Integração e Interoperabilidade")
replace_text(find_paragraph("4.2 Testes de Criptoagilidade"), "4.2 Testes de Criptoagilidade")
replace_text(find_paragraph("4.3 Desempenho e segurança"), "4.3 Desempenho e Segurança")


# Descrições das fases
insert_after(
    find_paragraph("1. Benchmarking e Informacional"),
    "Consolidar referências, requisitos e arquiteturas relevantes para orientar as decisões das fases seguintes e nivelar o conhecimento das equipes.",
)
insert_after(
    find_paragraph("2. Conceitual"),
    "Definir a arquitetura-alvo, a governança e o planejamento da evolução criptográfica antes da implementação.",
)
insert_after(
    find_paragraph("3. Desenvolvimento"),
    "Implementar as correções e os componentes definidos na fase conceitual, inicialmente em laboratório e ambiente de integração.",
)
insert_after(
    find_paragraph("4. Testes"),
    "Comprovar funcionamento, interoperabilidade, capacidade de transição, desempenho e segurança, incluindo cenários de falha e rollback.",
)


# 2. Conceitual
add_activity_content(
    "2.1 Definição e Detalhamento da Arquitetura",
    "Mapear a situação atual e definir a arquitetura-alvo e de transição para os ambientes de TI e tecnologia operacional das usinas.",
    [
        "Identificar ativos, aplicações, redes, fluxos críticos, protocolos e limites de confiança.",
        "Definir os componentes de proxy, gestão de certificados e chaves, inventário e monitoração.",
        "Definir requisitos de disponibilidade, latência, auditoria, recuperação e rollback.",
        "Validar a arquitetura com segurança, engenharia, operação e responsáveis pelas aplicações.",
    ],
    "3 a 5 semanas.",
)

add_activity_content(
    "2.2 Governança e Criptoagilidade",
    "Estabelecer a capacidade de substituir algoritmos, bibliotecas, certificados e provedores sem alterações extensas nas aplicações e sem interrupções desnecessárias.",
    [
        "Definir responsáveis, processo de decisão, exceções e comunicação das mudanças criptográficas.",
        "Criar uma política com mecanismos permitidos, transitórios e proibidos.",
        "Traduzir a política em perfis técnicos aplicáveis por configuração ou automação, quando possível.",
        "Definir indicadores de cobertura, conformidade, tempo, custo e facilidade de migração.",
        "Planejar exercícios periódicos de troca, coexistência e rollback.",
    ],
    "3 a 5 semanas.",
)

add_activity_content(
    "2.3 Inventário Lógico, Físico e Criptográfico",
    "Criar uma visão consolidada do uso de criptografia, relacionando dados críticos aos ativos e componentes que os protegem.",
    [
        "Inventariar algoritmos, protocolos, bibliotecas, certificados, chaves, hardware e firmware.",
        "Relacionar os itens aos dados protegidos, proprietários, criticidade, exposição e fornecedores.",
        "Registrar expiração, fim de suporte, capacidade de atualização e dificuldade de migração.",
        "Definir coleta recorrente e integração com os inventários corporativos existentes.",
    ],
    "4 a 7 semanas.",
)

add_activity_content(
    "2.4 Planejamento Criptográfico",
    "Transformar a política, o inventário e os riscos em um roadmap de correções, modernização e migração.",
    [
        "Priorizar ativos conforme criticidade dos dados, exposição e urgência da substituição.",
        "Definir ondas de migração, pilotos, critérios de entrada e saída e responsáveis.",
        "Planejar coexistência, rollback e controles compensatórios para ativos não atualizáveis.",
        "Incluir requisitos de criptoagilidade em aquisições, contratos e avaliações de fornecedores.",
        "Revisar periodicamente a política, o inventário, os riscos e o progresso das migrações.",
    ],
    "3 a 5 semanas.",
)


# 3. Desenvolvimento
add_activity_content(
    "3.1 Correções",
    "Corrigir configurações inseguras, dependências obsoletas e falhas no uso de certificados, chaves e protocolos.",
    [
        "Atualizar bibliotecas e provedores compatíveis com as aplicações.",
        "Remover mecanismos proibidos e fortalecer parâmetros e configurações.",
        "Corrigir validação, renovação, revogação e armazenamento de certificados e chaves.",
        "Adicionar testes de regressão e evidências antes e depois das mudanças.",
    ],
    "4 a 8 semanas por onda de aplicações.",
)

add_activity_content(
    "3.2 Proxy",
    "Implementar um proxy ou gateway para centralizar políticas criptográficas e modernizar fluxos de sistemas que não podem ser alterados diretamente.",
    [
        "Definir, para cada fluxo, terminação TLS, recriptografia ou encaminhamento transparente.",
        "Implementar a comunicação cliente-proxy-servidor, incluindo autenticação mútua quando necessária.",
        "Integrar emissão, renovação e rotação de certificados e chaves com PKI, KMS ou HSM.",
        "Aplicar perfis criptográficos centralizados e permitir a troca controlada de algoritmos e provedores.",
        "Implementar alta disponibilidade, health checks, comportamento de falha e rollback.",
        "Registrar handshake, algoritmo negociado, latência, erros, capacidade e expiração de certificados.",
    ],
    "6 a 10 semanas para a POC integrada.",
)

add_activity_content(
    "3.3 Desacoplamento das Aplicações da Criptografia",
    "Separar as regras de negócio das implementações criptográficas por meio de adaptadores, bibliotecas, APIs ou serviços controlados.",
    [
        "Mapear chamadas criptográficas e dependências incorporadas ao código.",
        "Definir uma interface simples, documentada e com padrões seguros.",
        "Implementar provedores intercambiáveis e seleção por configuração autorizada.",
        "Versionar formatos e metadados para permitir coexistência e leitura de dados antigos.",
        "Adaptar uma aplicação piloto e validar troca de provedor e rollback.",
    ],
    "6 a 10 semanas para a plataforma e a primeira aplicação piloto.",
)

add_activity_content(
    "3.4 QKD",
    "Integrar a distribuição quântica de chaves a um caso de uso controlado, mantendo autenticação, gestão de chaves e operação híbrida.",
    [
        "Montar ambiente simulado ou de laboratório com os nós e canais necessários.",
        "Integrar a interface QKD ao KMS, HSM ou gerenciador de chaves.",
        "Conectar o consumidor de chaves, como encryptor, VPN ou aplicação.",
        "Definir estoque, consumo, rotação, expiração e descarte das chaves.",
        "Implementar fallback, failback, alarmes e monitoração sem registrar material de chave.",
    ],
    "8 a 14 semanas em laboratório; a implantação física depende de aquisição e infraestrutura.",
)


# 4. Testes
add_activity_content(
    "4.1 Testes Funcionais, de Integração e Interoperabilidade",
    "Validar os fluxos ponta a ponta, as interfaces e a coexistência entre componentes atuais e novos.",
    [
        "Testar certificados, chaves, políticas, proxy, aplicações e integrações QKD.",
        "Validar clientes, servidores, bibliotecas, versões e equipamentos do piloto.",
        "Executar cenários nominais, negativos e de incompatibilidade.",
        "Registrar evidências, defeitos, correções e resultados de regressão.",
    ],
    "3 a 5 semanas.",
)

add_activity_content(
    "4.2 Testes de Criptoagilidade",
    "Comprovar que algoritmos, provedores, certificados ou perfis podem ser substituídos de forma controlada e mensurável.",
    [
        "Executar troca de perfil ou provedor e medir o tempo e o esforço necessários.",
        "Validar coexistência, leitura de dados existentes e retirada do mecanismo anterior.",
        "Testar rollback e retorno ao estado estável.",
        "Verificar proteção contra downgrade e seleção de mecanismos proibidos.",
        "Registrar lacunas e ações para melhorar a maturidade de criptoagilidade.",
    ],
    "3 a 5 semanas.",
)

add_activity_content(
    "4.3 Desempenho e Segurança",
    "Avaliar o impacto das mudanças e verificar a resiliência e a proteção dos componentes implementados.",
    [
        "Medir handshake, latência, vazão, CPU, memória e tamanho de certificados e mensagens.",
        "Executar carga e comparar os resultados com a linha de base.",
        "Simular falhas de proxy, PKI, KMS/HSM, comunicação e QKD.",
        "Validar alta disponibilidade, recuperação, fallback e failback.",
        "Revisar código, configurações, acessos, segredos, logs e superfícies expostas.",
    ],
    "4 a 6 semanas.",
)


# 5. Transferência de tecnologia
anchor = find_paragraph_with_style("5. Transferência de Tecnologia", "Heading 3")
anchor = insert_after(
    anchor,
    "Preparar as equipes para instalar, operar, monitorar, manter e evoluir a solução com autonomia.",
)
anchor = insert_label_after(anchor, "Atividades:")
for text in [
    "Consolidar arquitetura, interfaces, configurações e decisões técnicas.",
    "Produzir guias de instalação, atualização, certificados, chaves, proxy, QKD e rollback.",
    "Realizar treinamentos específicos para arquitetura, desenvolvimento, segurança e operação.",
    "Executar laboratórios de mudança de perfil, falha, recuperação e análise de logs.",
    "Realizar handover, operação assistida e avaliação da autonomia das equipes.",
]:
    anchor = insert_bullet_after(anchor, text)
anchor = insert_label_after(anchor, "Estimativa: ", "6 a 10 semanas, com sobreposição aos testes.")
anchor = insert_label_after(anchor, "Referência técnica: ", "NIST CSWP 39-upd1 - Considerations for Achieving Crypto Agility: Strategies and Practices.")


doc.core_properties.title = "Projeto Usinas de Geração de Energia"
doc.core_properties.subject = "Plano de atividades complementado com criptoagilidade e proxy"
doc.save(OUTPUT)
print(OUTPUT)
