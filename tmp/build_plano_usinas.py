from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\Ludmilla\Documents\WebsitePortifolio\outputs\Plano_de_Atividades_Projeto_Usinas_Energia.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

FONT = "Calibri"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D7DEE8"
WHITE = "FFFFFF"
INK = "1E252D"
GOLD = "B98900"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    for tag, attrs in (
        ("w:tblW", {"w:w": str(total), "w:type": "dxa"}),
        ("w:tblInd", {"w:w": str(indent_dxa), "w:type": "dxa"}),
        ("w:tblLayout", {"w:type": "fixed"}),
    ):
        node = tbl_pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tbl_pr.append(node)
        for key, value in attrs.items():
            node.set(qn(key), value)

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="C9D2DE", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph, field_name):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, fld_end])
    set_run_font(run, size=9, color=MUTED)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_left_border(paragraph, color=BLUE, size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element

    def next_id(tag, attr):
        values = [int(n.get(qn(attr))) for n in numbering.findall(qn(tag)) if n.get(qn(attr))]
        return max(values or [0]) + 1

    def make_abstract(num_fmt, level_text, left=540, hanging=270, marker_font=None):
        abstract_id = next_id("w:abstractNum", "w:abstractNumId")
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), level_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, fmt, text, jc, p_pr])
        if marker_font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), marker_font)
            fonts.set(qn("w:hAnsi"), marker_font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num_id = next_id("w:num", "w:numId")
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)
        return num_id

    return make_abstract("bullet", "•", marker_font="Symbol"), make_abstract("decimal", "%1.", left=540, hanging=270)


def apply_num(paragraph, num_id, ilvl=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl_el, num_id_el])


def add_bullet(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph(style="Plan Bullet")
    apply_num(p, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_number(doc, text, num_id):
    p = doc.add_paragraph(style="Plan Number")
    apply_num(p, num_id)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_label_paragraph(doc, label, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(label)
    set_run_font(r, bold=True, color=DARK_BLUE)
    r2 = p.add_run(text)
    set_run_font(r2, color=INK)
    return p


def add_estimate(doc, estimate, calendar_note=None):
    p = doc.add_paragraph(style="Estimate")
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, LIGHT_GRAY)
    set_left_border(p, color=BLUE)
    r = p.add_run("Estimativa: ")
    set_run_font(r, bold=True, color=NAVY)
    r2 = p.add_run(estimate)
    set_run_font(r2, bold=True, color=NAVY)
    if calendar_note:
        r3 = p.add_run(f" | {calendar_note}")
        set_run_font(r3, color=MUTED)
    return p


def add_activity(doc, aid, title, description, scope, deliverables, acceptance, estimate, dependencies=None, risks=None):
    h = doc.add_paragraph(style="Heading 2")
    h.paragraph_format.page_break_before = False
    r = h.add_run(f"{aid} {title}")
    set_run_font(r, size=13, color=BLUE, bold=False)
    p = doc.add_paragraph(description)
    p.paragraph_format.keep_together = True
    add_label_paragraph(doc, "Escopo. ", "")
    for item in scope:
        add_bullet(doc, item, BULLET_ID)
    add_label_paragraph(doc, "Entregáveis. ", "")
    for item in deliverables:
        add_bullet(doc, item, BULLET_ID)
    add_label_paragraph(doc, "Critério de conclusão. ", acceptance)
    if dependencies:
        add_label_paragraph(doc, "Dependências. ", dependencies)
    if risks:
        add_label_paragraph(doc, "Atenções. ", risks)
    add_estimate(doc, estimate)


def add_phase_intro(doc, phase, title, objective, duration, outputs):
    h = doc.add_paragraph(style="Heading 1")
    h.paragraph_format.page_break_before = True
    r = h.add_run(f"{phase}. {title}")
    set_run_font(r, size=16, color=BLUE, bold=False)
    p = doc.add_paragraph(objective)
    p.paragraph_format.keep_together = True
    add_estimate(doc, duration, "duração de calendário, considerando frentes paralelas")
    add_label_paragraph(doc, "Saídas principais. ", outputs)


def add_table(doc, headers, rows, widths_dxa, alignments=None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_run_font(r, size=9.2, bold=True, color=NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if alignments and idx < len(alignments):
                p.alignment = alignments[idx]
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = False
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for name in ("Plan Bullet", "Plan Number"):
    if name not in styles:
        st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = styles[name]
    st.base_style = normal
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    st.font.size = Pt(11)
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.25

if "Estimate" not in styles:
    estimate_style = styles.add_style("Estimate", WD_STYLE_TYPE.PARAGRAPH)
else:
    estimate_style = styles["Estimate"]
estimate_style.base_style = normal
estimate_style.font.name = FONT
estimate_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
estimate_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
estimate_style.font.size = Pt(10.5)
estimate_style.paragraph_format.left_indent = Inches(0.08)
estimate_style.paragraph_format.right_indent = Inches(0.04)
estimate_style.paragraph_format.space_before = Pt(5)
estimate_style.paragraph_format.space_after = Pt(10)
estimate_style.paragraph_format.line_spacing = 1.15

BULLET_ID, NUMBER_ID = create_numbering(doc)

first_header = section.first_page_header.paragraphs[0]
first_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
first_header.paragraph_format.space_after = Pt(0)
r = first_header.add_run("PLANO DE ATIVIDADES  |  VERSÃO 1.0")
set_run_font(r, size=8.5, color=MUTED, bold=True)

header = section.header.paragraphs[0]
header.paragraph_format.space_after = Pt(0)
r = header.add_run("PROJETO USINAS DE GERAÇÃO DE ENERGIA")
set_run_font(r, size=8.5, color=MUTED, bold=True)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.paragraph_format.space_before = Pt(0)
r = footer.add_run("Plano de Atividades  |  Página ")
set_run_font(r, size=9, color=MUTED)
add_page_field(footer, "PAGE")

# Capa
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(72)
kicker = doc.add_paragraph()
kicker.paragraph_format.space_after = Pt(8)
r = kicker.add_run("PLANEJAMENTO TÉCNICO E EXECUTIVO")
set_run_font(r, size=10, color=GOLD, bold=True)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(10)
title.paragraph_format.keep_with_next = True
r = title.add_run("Projeto Usinas de\nGeração de Energia")
set_run_font(r, size=30, color=NAVY, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(26)
r = subtitle.add_run("Plano de atividades para arquitetura criptográfica, criptoagilidade, integração QKD, testes e transferência de tecnologia")
set_run_font(r, size=13.5, color=DARK_BLUE)

cover_rows = [
    ("Escopo principal", "Fases 2 a 5: Conceitual, Desenvolvimento, Testes e Transferência de Tecnologia"),
    ("Contexto", "Ambientes de TI e TO de usinas de geração de energia"),
    ("Horizonte estimado", "42 a 56 semanas para POC/piloto; implantação física de QKD pode acrescentar 8 a 16 semanas"),
    ("Data", "Agosto de 2026"),
]
add_table(doc, ["Campo", "Definição"], cover_rows, [2100, 7260], font_size=10)

note = doc.add_paragraph(style="Estimate")
set_paragraph_shading(note, LIGHT_GRAY)
set_left_border(note, color=GOLD)
r = note.add_run("Referência utilizada. ")
set_run_font(r, bold=True, color=NAVY)
r = note.add_run("O PDF fornecido foi usado somente como referência de organização e nível de detalhamento. Seu conteúdo não foi tratado como instrução para este projeto.")
set_run_font(r, color=INK)

doc.add_page_break()

# 1. Orientação
h = doc.add_paragraph(style="Heading 1")
r = h.add_run("1. Objetivo, escopo e base das estimativas")
set_run_font(r, size=16, color=BLUE)
doc.add_paragraph(
    "Este documento define as atividades necessárias para conceber, desenvolver, testar e transferir uma solução de segurança criptográfica aplicável a usinas de geração de energia. O foco começa na fase Conceitual e cobre duas frentes coordenadas: modernização criptográfica, incluindo proxy/gateway, correções e desacoplamento das aplicações; e avaliação e integração de distribuição quântica de chaves (QKD)."
)

nist_note = doc.add_paragraph(style="Estimate")
set_paragraph_shading(nist_note, LIGHT_GRAY)
set_left_border(nist_note, color=BLUE)
r = nist_note.add_run("Referencial de criptoagilidade. ")
set_run_font(r, bold=True, color=NAVY)
r = nist_note.add_run("As atividades de governança, inventário, automação, priorização, arquitetura, cadeia de suprimentos, testes de transição e maturidade foram alinhadas ao NIST CSWP 39-upd1, Considerations for Achieving Crypto Agility: Strategies and Practices (atualizado em 29/06/2026).")
set_run_font(r, color=INK)

add_label_paragraph(doc, "Objetivo do projeto. ", "estabelecer uma arquitetura segura, criptoágil, auditável e compatível com as restrições de disponibilidade dos ambientes de tecnologia operacional (TO), reduzindo dependências rígidas de algoritmos, bibliotecas, certificados e mecanismos de gestão de chaves.")

h = doc.add_paragraph(style="Heading 2")
r = h.add_run("1.1 Premissas de planejamento")
set_run_font(r, size=13, color=BLUE)
premissas = [
    "As estimativas representam duração de calendário, não esforço individual. Elas pressupõem execução paralela por uma equipe multidisciplinar.",
    "O cenário-base considera uma POC e um piloto controlado, sem alteração direta de funções de proteção, controle ou segurança operacional da planta sem aprovação formal.",
    "O inventário, a arquitetura atual, os requisitos de disponibilidade e as janelas de manutenção ainda precisam ser confirmados com cada usina.",
    "A frente QKD pode iniciar em simulador ou laboratório. Aquisição de equipamentos, disponibilidade de fibra, distância, atenuação e logística de instalação podem ampliar o prazo.",
    "Toda migração deve prever coexistência, reversão e operação híbrida com mecanismos criptográficos convencionais aprovados.",
]
for item in premissas:
    add_bullet(doc, item, BULLET_ID)

add_estimate(doc, "Margem recomendada: adicionar 25% a 40% ao prazo se o inventário estiver incompleto, houver dependência forte de fabricantes ou faltar ambiente de laboratório.")

h = doc.add_paragraph(style="Heading 2")
r = h.add_run("1.2 Equipe de referência")
set_run_font(r, size=13, color=BLUE)
roles = [
    ("Gestão", "gerente de projeto/produto e responsável técnico"),
    ("Arquitetura", "arquiteto de solução, arquiteto de segurança e especialista em TO"),
    ("Engenharia", "desenvolvedores, redes, PKI/KMS/HSM, DevSecOps e integração"),
    ("QKD", "especialista em comunicações quânticas e fornecedor/laboratório, quando aplicável"),
    ("Qualidade", "QA, testes de segurança, desempenho e continuidade"),
    ("Operação", "representantes da usina, SOC/NOC, manutenção, infraestrutura e donos das aplicações"),
]
add_table(doc, ["Frente", "Perfis envolvidos"], roles, [1800, 7560], font_size=10)

h = doc.add_paragraph(style="Heading 2")
r = h.add_run("1.3 Visão executiva do cronograma")
set_run_font(r, size=13, color=BLUE)
summary_rows = [
    ("2", "Conceitual", "10-14 sem.", "Semanas 1-14", "Governança, arquitetura, inventário, risco e roadmap aprovados"),
    ("3", "Desenvolvimento", "16-24 sem.", "Semanas 11-36", "POCs e piloto de proxy, criptoagilidade e QKD"),
    ("4", "Testes", "8-14 sem.", "Semanas 29-44", "Evidências funcionais, segurança, desempenho e aceite"),
    ("5", "Transferência", "6-10 sem.", "Semanas 36-48", "Documentação, capacitação e operação assistida"),
]
add_table(doc, ["Fase", "Nome", "Duração", "Janela", "Resultado principal"], summary_rows, [650, 1800, 1250, 1600, 4060], font_size=9)
add_estimate(doc, "Prazo total base: 42 a 56 semanas (aproximadamente 10 a 14 meses).", "as fases se sobrepõem; os prazos não devem ser somados de forma linear")

# Fase 2
add_phase_intro(
    doc,
    "2",
    "Conceitual",
    "A fase Conceitual transforma as necessidades de segurança, disponibilidade e continuidade das usinas em uma arquitetura-alvo, requisitos verificáveis e um plano de implantação. Ela reduz incertezas antes do desenvolvimento e determina onde criptografia pós-quântica, criptoagilidade e QKD são tecnicamente justificáveis.",
    "10 a 14 semanas",
    "governança e política; arquitetura atual e alvo; inventário criptográfico; ferramentas; risco; cadeia de suprimentos; estratégia QKD; backlog e plano de POC.",
)

add_activity(
    doc, "2.1", "Definição e detalhamento da arquitetura",
    "Mapear a arquitetura atual e definir a arquitetura-alvo para os domínios de TI, TO e interconexões, deixando explícitos os limites de confiança, fluxos críticos e pontos de aplicação de controles criptográficos.",
    [
        "Identificar zonas, conduítes, centros de operação, datacenters, acesso remoto, integrações corporativas e enlaces entre usinas.",
        "Modelar fluxos entre aplicações, gateways, historiadores, SCADA/DCS, serviços de identidade, PKI, KMS/HSM e monitoração.",
        "Definir componentes lógicos: camada de abstração criptográfica, proxy/gateway, gestão de certificados, gestão de chaves, inventário e telemetria.",
        "Definir requisitos não funcionais: disponibilidade, latência, capacidade, rastreabilidade, recuperação, manutenção e segregação.",
        "Produzir arquitetura de transição com coexistência e reversão segura.",
    ],
    [
        "Diagramas de contexto, contêineres/componentes e implantação.",
        "Catálogo de interfaces, fluxos, protocolos e dependências.",
        "Documento de arquitetura atual, arquitetura-alvo e arquitetura de transição.",
        "Registro de decisões de arquitetura e riscos iniciais.",
    ],
    "arquitetura revisada por segurança, engenharia, operação da usina e donos das aplicações, com decisões críticas registradas e pendências atribuídas.",
    "3 a 5 semanas",
    "acesso aos diagramas existentes, entrevistas com equipes e definição preliminar do escopo das usinas.",
    "mudanças em sistemas de TO exigem análise de impacto operacional e janelas controladas.",
)

add_activity(
    doc, "2.2", "Governança, política criptográfica e modelo de criptoagilidade",
    "Integrar a criptoagilidade à governança e à gestão de riscos da organização, definindo responsabilidades, políticas e mecanismos para substituir algoritmos, bibliotecas, parâmetros, certificados e provedores sem interromper a operação.",
    [
        "Designar um responsável executivo/técnico e estabelecer RACI para decisões, implantação, exceções e monitoramento.",
        "Relacionar requisitos legais, normativos, internos e de missão aos dados e casos de uso prioritários.",
        "Definir política criptográfica com mecanismos permitidos, transicionais e proibidos, incluindo datas, exceções e responsáveis.",
        "Traduzir a política em perfis de configuração consumíveis por automação, sempre que tecnicamente possível.",
        "Classificar pontos de acoplamento criptográfico em código, configuração, biblioteca, protocolo, certificado e hardware.",
        "Definir interfaces estáveis para cifrar, assinar, verificar, negociar sessões e obter chaves.",
        "Estabelecer política de versões, seleção de algoritmos, feature flags, rotação e desativação controlada.",
        "Definir estratégia híbrida e compatibilidade regressiva durante a migração.",
        "Definir indicadores de tempo, custo, facilidade de migração, cobertura do inventário e aderência à política.",
    ],
    [
        "Política criptográfica e modelo de governança.",
        "RACI e fluxo de decisão/depreciação de algoritmos.",
        "Modelo de perfis de política aplicáveis por automação.",
        "Matriz aplicação x dependência criptográfica x estratégia de desacoplamento.",
        "Contrato de interfaces e requisitos de configuração.",
        "Painel inicial de indicadores e meta de maturidade.",
    ],
    "política e responsabilidades aprovadas; pelo menos dois cenários de troca simulados no desenho, com impacto, comunicação, reversão e indicadores definidos.",
    "3 a 5 semanas",
    "arquitetura preliminar e representantes das aplicações prioritárias.",
)

add_activity(
    doc, "2.3", "Inventário centrado em ativos e dados",
    "Construir uma linha de base confiável do uso de criptografia, partindo da criticidade dos dados e casos de uso e relacionando-os aos ativos, aplicações, enlaces, certificados, chaves, protocolos, bibliotecas, firmware e equipamentos.",
    [
        "Consolidar inventários de ativos de TI/TO, topologia física, sistemas, versões e proprietários.",
        "Descobrir certificados, autoridades certificadoras, datas de expiração, tamanhos de chave, algoritmos e usos.",
        "Identificar TLS, VPN, SSH, IPsec, assinatura de código, bancos de dados, backups, firmware e canais proprietários.",
        "Registrar local de armazenamento e proteção de chaves, incluindo arquivos, cofres, KMS e HSM.",
        "Relacionar dados em repouso, em trânsito e em uso aos controles criptográficos aplicados.",
        "Classificar criticidade, exposição, tempo de retenção dos dados, proprietário, dependência externa e dificuldade de migração.",
    ],
    [
        "Modelo de dados do inventário e dicionário de campos.",
        "Inventário inicial com proprietário, criticidade e evidência de descoberta.",
        "Mapa de calor de risco criptográfico e lacunas de cobertura.",
        "Plano de atualização contínua e integração com CMDB/gestão de ativos.",
    ],
    "ativos prioritários cobertos, dados validados por seus proprietários e lacunas registradas com plano de tratamento.",
    "4 a 7 semanas",
    "acesso controlado aos ambientes, exportações de CMDB e colaboração de fabricantes/terceiros.",
    "varreduras em TO devem ser passivas ou previamente homologadas para não afetar equipamentos sensíveis.",
)

add_activity(
    doc, "2.4", "Diagnóstico de ferramentas e automação",
    "Avaliar se as ferramentas corporativas conseguem descobrir, caracterizar, aplicar políticas e monitorar criptografia, identificando lacunas antes de escolher novas soluções.",
    [
        "Avaliar CMDB, gestão de ativos, vulnerabilidades, configuração, logs, repositórios, pipelines e inventário de software/hardware.",
        "Verificar capacidade de descobrir algoritmos, bibliotecas, protocolos, comprimentos de chave, certificados, dependências e firmware.",
        "Definir integrações e um modelo comum de dados para evitar inventários isolados.",
        "Selecionar oportunidades de política como código, validação em pipeline, bloqueio de configurações e alertas contínuos.",
        "Definir controles compensatórios quando um ativo não puder ser tornado criptoágil no prazo necessário.",
    ],
    [
        "Matriz de capacidade e lacunas das ferramentas.",
        "Arquitetura de integração e modelo de dados.",
        "Backlog de automação e política como código.",
        "Catálogo de controles compensatórios para ativos não ágeis.",
    ],
    "lacunas priorizadas por risco e esforço, com decisão de integrar, ampliar, substituir ou compensar cada capacidade crítica.",
    "2 a 4 semanas",
    "inventário inicial, acesso às ferramentas e participação de suas equipes responsáveis.",
)

add_activity(
    doc, "2.5", "Priorização de riscos e roadmap de migração",
    "Transformar riscos e requisitos em uma sequência priorizada de correções, modernizações e migrações, com critérios claros para algoritmos clássicos, pós-quânticos e uso de QKD.",
    [
        "Definir padrões permitidos, legados tolerados temporariamente e mecanismos a descontinuar.",
        "Priorizar casos considerando política, criticidade, exposição, vida útil do dado, dependência de fornecedor, tempo/custo de migração e facilidade de reversão.",
        "Definir ondas de migração, janelas de mudança, pilotos, indicadores e gates de aprovação.",
        "Estabelecer plano de renovação de certificados, chaves, bibliotecas e equipamentos.",
        "Definir governança, responsáveis, exceções e processo de resposta a vulnerabilidades criptográficas.",
    ],
    [
        "Roadmap de 12, 24 e 36 meses.",
        "Matriz de priorização e plano de ondas.",
        "Políticas técnicas mínimas e processo de exceção.",
        "Plano de reversão e continuidade por onda.",
    ],
    "roadmap aprovado pelos responsáveis de segurança, arquitetura, operação e orçamento, com prioridades e critérios de entrada/saída definidos.",
    "3 a 5 semanas",
    "inventário inicial e arquitetura-alvo suficientemente maduros.",
)

add_activity(
    doc, "2.6", "Cadeia de suprimentos e requisitos de aquisição",
    "Incorporar criptoagilidade às relações com fabricantes, integradores e fornecedores, reduzindo o risco de componentes que não possam receber novos algoritmos, parâmetros ou atualizações.",
    [
        "Levantar dependências de hardware, firmware, bibliotecas, protocolos, licenças, serviços e ciclos de suporte.",
        "Solicitar aos fornecedores uma lista dos componentes criptográficos e sua capacidade de atualização/substituição.",
        "Definir requisitos de aquisição: algoritmos configuráveis, atualização segura, exportação de inventário, prazos de correção, suporte a transições e evidências de validação.",
        "Classificar itens por fim de suporte, substituição obrigatória, mitigação temporária ou modernização programada.",
        "Criar cláusulas e critérios técnicos para POCs, contratos e homologação.",
    ],
    [
        "Matriz fornecedor x componente x capacidade de atualização.",
        "Questionário e requisitos de criptoagilidade para aquisição.",
        "Plano de tratamento de fim de vida/suporte.",
        "Riscos e dependências de cadeia de suprimentos registrados.",
    ],
    "fornecedores críticos avaliados e requisitos incorporados ao processo de compra, renovação e homologação.",
    "3 a 6 semanas",
    "lista de fornecedores, contratos, responsáveis por aquisição e acesso a informações técnicas.",
)

add_activity(
    doc, "2.7", "Arquitetura de APIs, gateways e protocolos criptoágeis",
    "Definir onde utilizar uma API criptográfica, provedores intercambiáveis, proxy/gateway ou mecanismos nativos de negociação, equilibrando modularidade, interoperabilidade e simplicidade operacional.",
    [
        "Definir identificadores e metadados de algoritmo/versão para dados, mensagens e protocolos.",
        "Especificar lista permitida e conjunto mínimo de interoperabilidade por caso de uso.",
        "Proteger a negociação contra downgrade e impedir retorno silencioso a mecanismos proibidos.",
        "Manter suporte coexistente somente pelo período necessário e remover opções obsoletas após a transição.",
        "Definir APIs simples, bem documentadas e com padrões seguros, evitando expor escolhas criptográficas desnecessárias à aplicação.",
        "Selecionar gateway para sistemas legados quando a alteração direta for inviável, com alta disponibilidade e controle de política centralizado.",
    ],
    [
        "Padrão de API e provedores.",
        "Perfis de negociação/interoperabilidade.",
        "Requisitos contra downgrade e de remoção de legado.",
        "Matriz de decisão API x gateway x alteração nativa.",
    ],
    "cada caso prioritário possui padrão de integração, conjunto permitido, estratégia de coexistência e condição de retirada do legado.",
    "3 a 5 semanas",
    "política criptográfica, arquitetura-alvo e inventário de interfaces.",
)

add_activity(
    doc, "2.8", "Concepção da arquitetura QKD",
    "Avaliar onde QKD agrega valor e desenhar uma arquitetura de laboratório/piloto integrada à gestão de chaves existente, sem presumir que QKD substitui autenticação, PKI ou criptografia de dados.",
    [
        "Selecionar enlaces candidatos com base em criticidade, distância, disponibilidade de fibra, atenuação e topologia.",
        "Definir nós QKD, canal quântico, canal clássico autenticado e pontos de confiança.",
        "Especificar integração com gerenciador de chaves, KMS/HSM, encryptors, VPN, TLS ou aplicação.",
        "Definir política de consumo de chaves, sincronização, buffer, expiração e descarte.",
        "Definir operação híbrida, fallback, alarmes e resposta à indisponibilidade do enlace QKD.",
        "Planejar homologação em simulador, laboratório e piloto físico.",
    ],
    [
        "Arquitetura QKD conceitual e de integração.",
        "Matriz de enlaces candidatos e critérios de viabilidade.",
        "Requisitos de interface e segurança para gestão de chaves.",
        "Plano de POC QKD e critérios de sucesso.",
    ],
    "um enlace candidato e um caso de uso aprovados para POC, com interfaces, fallback, indicadores e restrições documentados.",
    "3 a 5 semanas",
    "dados do enlace, participação de redes/telecom e acesso ao fornecedor ou laboratório QKD.",
    "QKD protege a distribuição de chaves; autenticação, proteção de endpoints e governança continuam necessárias.",
)

add_activity(
    doc, "2.9", "Planejamento da POC e backlog técnico",
    "Converter a arquitetura em histórias técnicas, experimentos controlados e critérios de aceite para as frentes de proxy, desacoplamento, inventário e QKD.",
    [
        "Definir ambientes, dados sintéticos, massas de teste, ferramentas e responsáveis.",
        "Especificar requisitos rastreáveis e critérios de aceite funcionais e não funcionais.",
        "Planejar sprints, dependências, riscos, aprovações e recursos de laboratório.",
        "Definir o que não será testado em produção e quais evidências são necessárias para avançar ao piloto.",
    ],
    [
        "Backlog priorizado e estimado.",
        "Plano de POC/piloto e matriz de rastreabilidade.",
        "Plano de ambientes, dados, ferramentas e acessos.",
        "Critérios de entrada, saída e interrupção segura.",
    ],
    "backlog pronto para execução, dependências críticas encaminhadas e critérios de aceite aprovados.",
    "2 a 3 semanas",
    "decisões de arquitetura e escopo do piloto.",
)

# Fase 3
add_phase_intro(
    doc,
    "3",
    "Desenvolvimento",
    "A fase de Desenvolvimento implementa os componentes e adaptações definidos na arquitetura. O trabalho deve ocorrer primeiro em laboratório e ambientes de integração, com automação, versionamento, evidências e mecanismos de reversão desde o início.",
    "16 a 24 semanas",
    "ambientes reproduzíveis; correções criptográficas; proxy/gateway; camada de criptoagilidade; inventário automatizado; integração QKD; telemetria e documentação técnica.",
)

add_activity(
    doc, "3.1", "Ambiente de engenharia e linha de base",
    "Preparar um ambiente reproduzível para compilar, configurar, testar e comparar os componentes sem dependência do ambiente produtivo.",
    [
        "Provisionar repositórios, pipelines, gestão de segredos, registros de artefatos e controle de versões.",
        "Montar topologia de referência com clientes, proxy/gateway, serviços, PKI/KMS/HSM simulados e observabilidade.",
        "Registrar métricas da solução atual para comparação.",
        "Definir padrões de logs, rastreabilidade e evidências de teste.",
    ],
    ["Ambiente de laboratório automatizado.", "Pipeline de build/deploy/teste.", "Baseline funcional e de desempenho.", "Guia de reprodução do ambiente."],
    "um novo ambiente pode ser criado a partir do repositório e executar o teste básico ponta a ponta sem configuração manual não documentada.",
    "2 a 4 semanas",
    "infraestrutura, acessos e ferramentas aprovadas.",
)

add_activity(
    doc, "3.2", "Correções e modernização criptográfica",
    "Corrigir vulnerabilidades, configurações inseguras e dependências obsoletas identificadas no inventário, priorizando mudanças de baixo risco e alto impacto.",
    [
        "Atualizar bibliotecas e provedores criptográficos compatíveis com a aplicação.",
        "Remover algoritmos, protocolos e parâmetros proibidos conforme política aprovada.",
        "Corrigir validação de certificados, cadeia de confiança, nomes, expiração e revogação.",
        "Eliminar chaves ou segredos fixos em código/configuração e integrar armazenamento seguro.",
        "Adicionar testes de regressão e verificação de configuração.",
    ],
    ["Pacotes de correção por aplicação.", "Evidências antes/depois.", "Testes automatizados.", "Registro de exceções e itens dependentes de fornecedor."],
    "correções críticas aprovadas e testadas, sem regressão funcional e com plano de reversão validado.",
    "4 a 8 semanas por onda de 3 a 5 aplicações",
    "priorização do inventário, ambientes de teste e disponibilidade dos donos das aplicações.",
)

add_activity(
    doc, "3.3", "Proxy/Gateway criptográfico",
    "Implementar um proxy ou gateway para centralizar políticas criptográficas, terminação ou encaminhamento de sessões, gestão de certificados e coleta de métricas, reduzindo mudanças imediatas nas aplicações legadas.",
    [
        "Definir modo de operação por fluxo: terminação, recriptografia, passagem transparente ou integração lateral.",
        "Configurar TLS mútuo quando aplicável, cadeia de confiança, renovação e rotação de certificados.",
        "Implementar políticas por serviço, segregação, limites, health checks e proteção contra falhas em cascata.",
        "Adicionar telemetria de handshake, erros, latência, algoritmo negociado e expiração.",
        "Validar compatibilidade clássica e preparar perfis híbridos/pós-quânticos no ambiente experimental.",
        "Aplicar política centralizada, rotação e novos algoritmos de modo uniforme quando os componentes legados não puderem ser alterados diretamente.",
    ],
    ["Proxy/gateway implantável por automação.", "Perfis de configuração versionados.", "Integração com PKI/KMS e observabilidade.", "Runbook de operação e reversão."],
    "comunicação ponta a ponta estável, políticas aplicadas, certificados validados, métricas coletadas e fallback testado.",
    "6 a 10 semanas",
    "linha de base, definição dos fluxos-piloto e certificados de laboratório.",
    "o proxy não deve criar ponto único de falha; alta disponibilidade e capacidade precisam ser verificadas.",
)

add_activity(
    doc, "3.4", "Plano de mudança para desacoplar aplicações da criptografia",
    "Detalhar e preparar a sequência de refatoração que remove decisões criptográficas de regras de negócio e as transfere para interfaces, serviços ou adaptadores controlados.",
    [
        "Mapear chamadas criptográficas e dependências por módulo.",
        "Definir estratégia por aplicação: adaptador interno, biblioteca corporativa, serviço criptográfico ou proxy externo.",
        "Planejar compatibilidade de dados, certificados, formatos de assinatura e migração gradual.",
        "Criar etapas pequenas com feature flags, dupla leitura/validação, observabilidade e reversão.",
        "Estimar esforço por aplicação e organizar ondas de mudança.",
    ],
    ["Plano técnico por aplicação.", "Sequência de releases e critérios de promoção.", "Matriz de compatibilidade de dados/protocolos.", "Plano de rollback e comunicação de mudança."],
    "uma aplicação representativa possui plano executável, tarefas estimadas, dependências identificadas e rollback testável.",
    "3 a 6 semanas",
    "código-fonte, especialistas da aplicação e modelo de criptoagilidade.",
)

add_activity(
    doc, "3.5", "Camada de abstração e criptoagilidade",
    "Construir a interface configurável que permite trocar algoritmos, provedores e perfis sem alterar a lógica de negócio.",
    [
        "Implementar contratos de API/SDK para cifragem, assinatura, verificação e obtenção de chaves.",
        "Separar política criptográfica de implementação e permitir seleção por configuração autorizada.",
        "Projetar padrões seguros e uma interface simples e bem documentada, reduzindo escolhas desnecessárias para os desenvolvedores.",
        "Implementar versionamento de envelopes, metadados de algoritmo e compatibilidade de leitura.",
        "Integrar provedores clássicos e um provedor experimental adicional para demonstrar troca controlada.",
        "Adicionar auditoria, métricas, testes de contrato e validação de configuração.",
    ],
    ["SDK/serviço de abstração.", "Perfis de política versionados.", "Adaptador para aplicação piloto.", "Testes de troca, coexistência e rollback."],
    "troca de provedor/perfil demonstrada em laboratório sem alteração da regra de negócio, preservando leitura de dados existentes e capacidade de reversão.",
    "6 a 10 semanas",
    "contratos definidos e aplicação piloto disponível.",
)

add_activity(
    doc, "3.6", "Automação do inventário criptográfico",
    "Implementar coleta recorrente e consolidação do inventário para reduzir dependência de levantamentos manuais e apoiar alertas e priorização.",
    [
        "Criar conectores para certificados, endpoints, repositórios, pipelines, cofres, KMS/HSM e fontes de ativos aprovadas.",
        "Normalizar evidências, eliminar duplicidades e relacionar ativo, proprietário, serviço e dependência.",
        "Gerar alertas de expiração, algoritmo proibido, chave fraca, certificado desconhecido e lacuna de proprietário.",
        "Integrar perfis de política consumíveis por máquina para avaliar conformidade e recomendar ações.",
        "Publicar painéis e exportações para governança e auditoria.",
    ],
    ["Coletores automatizados.", "Base consolidada e versionada.", "Painel de risco e expiração.", "Procedimento de operação e tratamento de falsos positivos."],
    "execução recorrente produz resultados rastreáveis, com cobertura e taxa de erro conhecidas e processo de correção definido.",
    "4 a 7 semanas",
    "fontes de dados, credenciais de leitura e política de varredura em TO.",
)

add_activity(
    doc, "3.7", "Integração com PKI, KMS e HSM",
    "Integrar emissão, armazenamento, rotação, revogação e auditoria de certificados e chaves aos componentes desenvolvidos.",
    [
        "Automatizar solicitação, aprovação, emissão e renovação de certificados.",
        "Definir proteção e separação de chaves por ambiente, aplicação e criticidade.",
        "Integrar KMS/HSM ou simuladores com controle de acesso, logs e segregação de funções.",
        "Testar expiração, revogação, indisponibilidade e recuperação.",
    ],
    ["Fluxos automatizados de ciclo de vida.", "Políticas e perfis de certificado.", "Integração auditável com KMS/HSM.", "Runbooks de rotação, revogação e recuperação."],
    "ciclo de vida completo executado em laboratório, incluindo rotação e revogação, sem exposição de material sensível.",
    "5 a 8 semanas",
    "PKI/KMS/HSM disponíveis ou simuladores homologados.",
)

h = doc.add_paragraph(style="Heading 2")
r = h.add_run("3.8 Trilha QKD")
set_run_font(r, size=13, color=BLUE)
doc.add_paragraph("A trilha QKD deve ser desenvolvida como integração de gestão de chaves e não como substituição automática da criptografia de aplicações. O piloto deve manter autenticação forte, proteção de endpoints e um caminho híbrido de continuidade.")

qkd_activities = [
    ("3.8.1", "Ambiente QKD simulado ou de laboratório", "Montar os pares de nós, o canal clássico autenticado, o simulador/equipamento e as APIs de entrega de chaves.", ["Topologia reproduzível.", "Conectividade e sincronização validadas.", "Métricas de geração/estoque de chaves."], "4 a 7 semanas"),
    ("3.8.2", "Adaptador de gerência de chaves QKD", "Implementar um adaptador entre a interface do sistema QKD e o KMS/gerenciador de chaves, com autenticação, autorização, metadados, consumo e descarte seguros.", ["Adaptador versionado.", "Controle de acesso e auditoria.", "Testes de contrato e erros."], "6 a 10 semanas"),
    ("3.8.3", "Integração com encryptor, VPN, TLS ou aplicação", "Consumir chaves disponibilizadas pelo adaptador no caso de uso selecionado, mantendo separação entre distribuição de chaves e mecanismo de proteção de dados.", ["Integração ponta a ponta.", "Política de consumo/rotação.", "Evidências de confidencialidade e continuidade."], "6 a 10 semanas"),
    ("3.8.4", "Operação híbrida e fallback", "Implementar regras para perda do enlace QKD, baixo estoque de chaves, ressincronização e retorno controlado, usando mecanismo alternativo aprovado.", ["Máquina de estados de operação.", "Alarmes e limiares.", "Runbook e testes de failover/failback."], "4 a 7 semanas"),
    ("3.8.5", "Observabilidade QKD", "Coletar taxa de geração, estoque, consumo, erros, disponibilidade, qualidade do enlace e eventos de segurança, sem registrar material de chave.", ["Painel e alertas.", "Logs auditáveis.", "Indicadores para aceite do piloto."], "3 a 5 semanas"),
]
for aid, title_text, desc, entreg, est in qkd_activities:
    add_activity(
        doc, aid, title_text, desc,
        [
            "Implementar em laboratório com interfaces e configurações versionadas.",
            "Automatizar testes básicos, tratamento de erro e coleta de evidências.",
            "Documentar limites, pressupostos e comportamento de reversão.",
        ],
        entreg,
        "o cenário nominal e pelo menos um cenário de falha são reproduzíveis e possuem evidências suficientes para a fase de testes.",
        est,
        "arquitetura QKD aprovada, acesso ao simulador/equipamento e caso de uso selecionado.",
    )

add_estimate(doc, "Implantação física QKD: acrescentar tipicamente 8 a 16 semanas", "dependendo de aquisição, fibra, obras, licenciamento, logística e disponibilidade do fornecedor")

# Fase 4
add_phase_intro(
    doc,
    "4",
    "Testes e validação",
    "A fase de Testes comprova que as mudanças atendem aos requisitos sem degradar disponibilidade, desempenho ou segurança. Os testes devem cobrir operação nominal, falhas, regressão, coexistência e reversão, com evidências rastreáveis aos requisitos.",
    "8 a 14 semanas",
    "plano e evidências de teste; relatório de vulnerabilidades; baseline comparativa; resultados QKD; aceite técnico e operacional do piloto.",
)

test_activities = [
    ("4.1", "Planejamento, ambientes e rastreabilidade", "Consolidar cenários, massas, pré-condições, ferramentas, responsáveis e critérios de aprovação, ligando cada teste a requisitos e riscos.", "2 a 3 semanas"),
    ("4.2", "Testes funcionais e de integração", "Validar fluxos ponta a ponta, certificados, chaves, políticas, proxy, abstração criptográfica, inventário, APIs QKD e telemetria.", "3 a 5 semanas"),
    ("4.3", "Interoperabilidade e compatibilidade", "Testar versões, clientes, servidores, bibliotecas, dispositivos e fornecedores, incluindo coexistência com ambientes legados.", "3 a 5 semanas"),
    ("4.4", "Desempenho e capacidade", "Comparar handshake, latência, vazão, uso de CPU/memória, tamanhos, filas, taxa de chaves QKD e comportamento sob carga.", "3 a 6 semanas"),
    ("4.5", "Resiliência, continuidade e reversão", "Simular indisponibilidade de proxy, PKI, KMS/HSM, nó QKD, canal clássico, perda de enlace, esgotamento de chaves e recuperação.", "3 a 5 semanas"),
    ("4.6", "Segurança e hardening", "Executar revisão de configuração e código, análise de vulnerabilidades, testes de acesso, proteção de segredos, logs e superfícies expostas.", "4 a 6 semanas"),
    ("4.7", "Validação específica de QKD", "Medir disponibilidade do enlace, taxa/estoque de chaves, sincronização, consumo, descarte, alarmes, failover/failback e integração com o mecanismo de cifragem.", "3 a 6 semanas"),
    ("4.8", "Piloto e aceite em ambiente representativo", "Executar operação controlada com usuários técnicos e operação, acompanhar indicadores, corrigir defeitos e formalizar aceite ou pendências.", "3 a 6 semanas"),
]
for aid, title_text, desc, est in test_activities:
    add_activity(
        doc, aid, title_text, desc,
        [
            "Definir cenário nominal, cenários negativos e critérios quantitativos.",
            "Executar, registrar evidências, abrir defeitos e repetir após correções.",
            "Comparar resultados com a linha de base e documentar limitações.",
        ],
        ["Casos e scripts de teste.", "Evidências e métricas.", "Relatório de defeitos/riscos.", "Recomendação de aceite ou ajustes."],
        "todos os casos críticos aprovados; defeitos residuais classificados, aceitos formalmente e com responsável/prazo; rollback comprovado.",
        est,
        "componentes integrados, dados de teste e ambientes estáveis.",
    )

h = doc.add_paragraph(style="Heading 2")
r = h.add_run("4.9 Critérios mínimos de aceite do piloto")
set_run_font(r, size=13, color=BLUE)
acceptance_items = [
    "Nenhum impacto não autorizado em funções de proteção, controle ou disponibilidade da planta.",
    "Fluxos prioritários operam com autenticação, confidencialidade e integridade conforme a política aprovada.",
    "Troca de perfil/provedor e retorno ao estado anterior são reproduzíveis e auditáveis.",
    "Negociação de protocolos resiste a downgrade, não seleciona opções proibidas e mantém o conjunto mínimo de interoperabilidade aprovado.",
    "Proxy/gateway atende metas de capacidade e não representa ponto único de falha.",
    "Eventos, falhas e expirações geram telemetria acionável sem exposição de chaves ou segredos.",
    "No QKD, indisponibilidade ou estoque insuficiente aciona o comportamento híbrido definido, sem interrupção inesperada.",
    "Documentação, evidências, riscos residuais e pendências foram aceitos pelos responsáveis técnicos e operacionais.",
]
for item in acceptance_items:
    add_bullet(doc, item, BULLET_ID)

# Fase 5
add_phase_intro(
    doc,
    "5",
    "Transferência de tecnologia",
    "A transferência de tecnologia prepara as equipes para instalar, operar, monitorar, manter e evoluir a solução sem dependência excessiva da equipe de desenvolvimento ou de fornecedores.",
    "6 a 10 semanas",
    "documentação final; trilhas de capacitação; laboratórios guiados; runbooks; operação assistida; aceite de conhecimento e plano de evolução.",
)

transfer_activities = [
    ("5.1", "Documentação de arquitetura, engenharia e operação", "Consolidar diagramas, decisões, interfaces, configurações, segurança, procedimentos de instalação, manutenção, troubleshooting e recuperação.", "3 a 5 semanas"),
    ("5.2", "Capacitação por perfil", "Preparar módulos específicos para arquitetura, desenvolvimento, redes/infraestrutura, segurança/SOC, operação TO, suporte e gestão.", "2 a 4 semanas"),
    ("5.3", "Laboratórios e exercícios práticos", "Executar instalação, rotação de certificados, mudança de perfil, falha de proxy/KMS/QKD, fallback, recuperação e análise de logs em ambiente seguro.", "2 a 4 semanas"),
    ("5.4", "Handover operacional e governança", "Transferir rotinas, acessos, dashboards, SLAs, escalonamento, gestão de mudanças, exceções e relacionamento com fornecedores.", "3 a 6 semanas"),
    ("5.5", "Operação assistida e encerramento", "Acompanhar ciclos reais controlados, resolver dúvidas, medir autonomia, tratar pendências e formalizar aceite final.", "4 a 8 semanas"),
]
for aid, title_text, desc, est in transfer_activities:
    add_activity(
        doc, aid, title_text, desc,
        [
            "Definir público, pré-requisitos, objetivos e critérios de aprendizagem.",
            "Produzir material versionado e reutilizável.",
            "Executar atividades práticas e coletar evidências de autonomia.",
        ],
        ["Documentos e runbooks aprovados.", "Material de treinamento.", "Gravações ou roteiros de laboratório.", "Lista de presença, avaliação e plano de lacunas."],
        "equipes designadas executam os procedimentos críticos com sucesso, conhecem os limites da solução e sabem acionar suporte/escalonamento.",
        est,
        "solução estável, documentação preliminar e disponibilidade das equipes de destino.",
    )

# Integração e marcos
add_phase_intro(
    doc,
    "6",
    "Cronograma integrado, marcos e dependências",
    "O cronograma abaixo organiza as fases em ondas sobrepostas. O início de desenvolvimento pode ocorrer quando a arquitetura mínima e o backlog da POC estiverem aprovados; testes começam por componente antes da integração completa; documentação e capacitação evoluem durante todo o projeto.",
    "42 a 56 semanas para o cenário-base",
    "marcos executivos, gates de decisão, responsáveis e uma visão comum das dependências críticas.",
)

milestones = [
    ("M1", "Sem. 4", "Arquitetura atual e requisitos prioritários validados", "Autorizar detalhamento e POC"),
    ("M2", "Sem. 10-14", "Política, arquitetura-alvo, inventário, riscos e roadmap aprovados", "Autorizar desenvolvimento integrado"),
    ("M3", "Sem. 20-28", "Proxy/gateway e primeira aplicação criptoágil em laboratório", "Autorizar testes integrados"),
    ("M4", "Sem. 24-34", "Integração QKD funcional em simulador/laboratório", "Decidir piloto físico"),
    ("M5", "Sem. 38-44", "Piloto técnico e operacional concluído", "Aceite ou plano de ajustes"),
    ("M6", "Sem. 44-48", "Handover e capacitação concluídos", "Entrada em operação assistida"),
    ("M7", "Sem. 48-56", "Operação assistida e encerramento", "Aceite final e roadmap de evolução"),
]
add_table(doc, ["Marco", "Janela", "Evidência", "Decisão"], milestones, [900, 1250, 4510, 2700], font_size=9.2)

h = doc.add_paragraph(style="Heading 2")
r = h.add_run("6.1 Dependências críticas")
set_run_font(r, size=13, color=BLUE)
dependencies = [
    "Disponibilidade de diagramas, inventários, código-fonte e especialistas das aplicações.",
    "Ambiente de laboratório representativo e processo de acesso seguro aos ambientes de TO.",
    "Aprovação de políticas, perfis criptográficos e autoridade para emissão de certificados de teste.",
    "Compatibilidade e suporte de fabricantes de equipamentos e aplicações legadas.",
    "Definição do enlace QKD, fibra disponível, orçamento, prazo de aquisição e suporte do fornecedor/laboratório.",
    "Janelas de mudança, critérios de interrupção e responsáveis por rollback.",
]
for item in dependencies:
    add_bullet(doc, item, BULLET_ID)

h = doc.add_paragraph(style="Heading 2")
r = h.add_run("6.2 Decisões a obter no início do projeto")
set_run_font(r, size=13, color=BLUE)
questions = [
    "Quais usinas, enlaces, aplicações e ambientes estão dentro da primeira onda?",
    "Quais fluxos são críticos e quais limites de latência, disponibilidade e recuperação devem ser atendidos?",
    "Quais componentes podem receber proxy/gateway e quais exigem alteração de aplicação ou apoio do fabricante?",
    "Qual é a política criptográfica atual e quais mecanismos devem ser priorizados para correção ou migração?",
    "O objetivo da primeira entrega é POC, piloto operacional ou implantação produtiva?",
    "Qual caso de uso QKD tem enlace, valor, risco e viabilidade suficientes para o piloto?",
    "Qual mecanismo híbrido será usado quando QKD estiver indisponível ou sem estoque de chaves?",
    "Quais evidências, normas internas e aprovações são necessárias para aceite?",
]
for item in questions:
    add_bullet(doc, item, BULLET_ID)

# Apêndice consolidado
add_phase_intro(
    doc,
    "Apêndice A",
    "Matriz consolidada de atividades e estimativas",
    "A matriz resume as durações para facilitar a criação de cronogramas, termos de referência e planos de trabalho. As durações são faixas preliminares e devem ser refinadas após a validação do inventário e das dependências.",
    "uso para planejamento inicial",
    "lista única de atividades, prazos e dependências imediatas.",
)

matrix_rows = [
    ("2.1", "Arquitetura atual, alvo e de transição", "3-5 sem.", "Diagramas e entrevistas"),
    ("2.2", "Governança, política e criptoagilidade", "3-5 sem.", "Arquitetura e partes interessadas"),
    ("2.3", "Inventário centrado em ativos e dados", "4-7 sem.", "Acessos e fontes de ativos"),
    ("2.4", "Diagnóstico de ferramentas e automação", "2-4 sem.", "Inventário e equipes de ferramentas"),
    ("2.5", "Priorização de riscos e roadmap", "3-5 sem.", "Inventário e política"),
    ("2.6", "Cadeia de suprimentos e aquisições", "3-6 sem.", "Fornecedores e contratos"),
    ("2.7", "APIs, gateways e protocolos criptoágeis", "3-5 sem.", "Política e interfaces"),
    ("2.8", "Arquitetura QKD", "3-5 sem.", "Dados do enlace e fornecedor"),
    ("2.9", "Plano de POC e backlog", "2-3 sem.", "Escopo e decisões de arquitetura"),
    ("3.1", "Ambiente e linha de base", "2-4 sem.", "Infraestrutura e ferramentas"),
    ("3.2", "Correções criptográficas", "4-8 sem./onda", "Aplicações e testes"),
    ("3.3", "Proxy/gateway criptográfico", "6-10 sem.", "Fluxos-piloto e PKI"),
    ("3.4", "Plano de desacoplamento", "3-6 sem.", "Código e especialistas"),
    ("3.5", "Camada de criptoagilidade", "6-10 sem.", "Contratos e aplicação piloto"),
    ("3.6", "Automação do inventário", "4-7 sem.", "Fontes e credenciais de leitura"),
    ("3.7", "Integração PKI/KMS/HSM", "5-8 sem.", "Serviços ou simuladores"),
    ("3.8.1", "Ambiente QKD", "4-7 sem.", "Simulador/equipamento"),
    ("3.8.2", "Adaptador de chaves QKD", "6-10 sem.", "APIs e KMS"),
    ("3.8.3", "Integração QKD ao caso de uso", "6-10 sem.", "Adaptador e encryptor/aplicação"),
    ("3.8.4", "Operação híbrida e fallback", "4-7 sem.", "Política de continuidade"),
    ("3.8.5", "Observabilidade QKD", "3-5 sem.", "Métricas e integração"),
    ("4.1-4.8", "Testes e piloto", "8-14 sem.", "Componentes integrados"),
    ("5.1-5.5", "Transferência e operação assistida", "6-10 sem.", "Solução estável e equipes"),
]
add_table(doc, ["ID", "Atividade", "Estimativa", "Dependência principal"], matrix_rows, [850, 3840, 1500, 3170], font_size=8.8)

h = doc.add_paragraph(style="Heading 2")
r = h.add_run("Como refinar as estimativas")
set_run_font(r, size=13, color=BLUE)
refinement = [
    "Confirmar quantidade de usinas, aplicações, enlaces, certificados e fornecedores.",
    "Separar esforço por onda e classificar complexidade baixa, média ou alta.",
    "Identificar atividades que podem ocorrer em paralelo e especialistas compartilhados.",
    "Adicionar prazos de aquisição, contratação, logística e homologação de equipamentos.",
    "Aplicar reserva para janelas operacionais, aprovações, testes de regressão e correções.",
    "Reestimar ao final da fase Conceitual e após a primeira POC integrada.",
]
for item in refinement:
    add_bullet(doc, item, BULLET_ID)

# Apêndice NIST
add_phase_intro(
    doc,
    "Apêndice B",
    "Rastreabilidade com o NIST CSWP 39-upd1",
    "A tabela mostra como as principais orientações do documento Considerations for Achieving Crypto Agility: Strategies and Practices foram convertidas em atividades verificáveis neste plano. O alinhamento é uma interpretação aplicada ao contexto de usinas e não representa certificação ou endosso do NIST.",
    "referência de alinhamento",
    "mapeamento entre temas do NIST, atividades do projeto e evidências esperadas.",
)

nist_rows = [
    ("Governança e política", "2.2, 2.5, 6", "Política, RACI, riscos, indicadores e gates"),
    ("Inventário centrado em ativos/dados", "2.3, 3.6", "Inventário rastreável e cobertura conhecida"),
    ("Ferramentas e descoberta automatizada", "2.4, 3.6", "Lacunas, conectores, alertas e política consumível por máquina"),
    ("Priorização e mitigação contínua", "2.5, 4, 5.5", "Roadmap, riscos residuais, testes recorrentes e operação assistida"),
    ("Cadeia de suprimentos", "2.6, 6.1", "Requisitos de aquisição, capacidade de atualização e fim de vida"),
    ("Arquitetura criptográfica", "2.1, 2.7, 3.7", "Arquitetura de dados em repouso/trânsito/uso e gestão de chaves"),
    ("Crypto API e provedores", "2.7, 3.5", "Interface simples, padrões seguros e provedores intercambiáveis"),
    ("Gateway para legado", "2.7, 3.3", "Política centralizada e modernização sem alteração imediata do legado"),
    ("Protocolos e interoperabilidade", "2.7, 4.3, 4.6", "Perfis, proteção contra downgrade e retirada de opções obsoletas"),
    ("Recursos e desempenho", "2.1, 4.4", "Limites de CPU, memória, rede, tamanho e capacidade medidos"),
    ("Maturidade e capacitação", "2.2, 5", "Meta de maturidade, indicadores, treinamento e autonomia"),
]
add_table(doc, ["Tema do NIST", "Atividades", "Evidência no projeto"], nist_rows, [2600, 1700, 5060], font_size=8.8)

h = doc.add_paragraph(style="Heading 2")
r = h.add_run("Referência")
set_run_font(r, size=13, color=BLUE)
doc.add_paragraph(
    "National Institute of Standards and Technology (NIST). Considerations for Achieving Crypto Agility: Strategies and Practices. NIST CSWP 39-upd1, incluindo atualizações de 29 de junho de 2026. DOI: 10.6028/NIST.CSWP.39-upd1. Disponível em: https://nvlpubs.nist.gov/nistpubs/CSWP/NIST.CSWP.39-upd1.pdf."
)

doc.core_properties.title = "Plano de Atividades - Projeto Usinas de Geração de Energia"
doc.core_properties.subject = "Arquitetura criptográfica, criptoagilidade, QKD, testes e transferência de tecnologia"
doc.core_properties.author = "Equipe do Projeto"
doc.core_properties.keywords = "usinas, energia, criptografia, criptoagilidade, QKD, proxy, testes, transferência de tecnologia"

doc.save(OUT)
print(OUT)
