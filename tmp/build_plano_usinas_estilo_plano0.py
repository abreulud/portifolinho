from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


OUTPUT = Path(r"C:\Users\Ludmilla\Documents\WebsitePortifolio\outputs\Plano_Projeto_Usinas_Energia_Estilo_Plano_0.docx")
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

FONT = "Arial"
BLACK = "000000"
HEADING = "595959"


def set_font(run, size=11, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    rpr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def create_numbering(document):
    numbering = document.part.numbering_part.element

    def next_id(tag, attr):
        values = [int(node.get(qn(attr))) for node in numbering.findall(qn(tag)) if node.get(qn(attr))]
        return max(values or [0]) + 1

    def add_definition(fmt, text, left=720, hanging=360, marker_font=None):
        abstract_id = next_id("w:abstractNum", "w:abstractNumId")
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))

        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)

        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")

        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), str(left))
        indent.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "276")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, indent, spacing])
        level.extend([start, num_fmt, lvl_text, lvl_jc, ppr])

        if marker_font:
            rpr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), marker_font)
            fonts.set(qn("w:hAnsi"), marker_font)
            rpr.append(fonts)
            level.append(rpr)

        abstract.append(level)
        numbering.append(abstract)

        num_id = next_id("w:num", "w:numId")
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abstract_id))
        num.append(ref)
        numbering.append(num)
        return num_id

    return add_definition("bullet", "•", marker_font="Symbol"), add_definition("decimal", "%1.")


def apply_number(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, nid])
    ppr.insert(0, numpr)


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="Plano Bullet")
    apply_number(paragraph, BULLET_ID)
    set_font(paragraph.add_run(text))
    return paragraph


def add_number(document, text):
    paragraph = document.add_paragraph(style="Plano Number")
    apply_number(paragraph, NUMBER_ID)
    set_font(paragraph.add_run(text))
    return paragraph


def add_label(document, label, text="", after=5):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = not bool(text)
    set_font(paragraph.add_run(label), bold=True)
    if text:
        set_font(paragraph.add_run(text))
    return paragraph


def add_component(document, name, description):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(7)
    set_font(paragraph.add_run(name), bold=True)
    paragraph = document.add_paragraph(description)
    paragraph.paragraph_format.space_after = Pt(8)


def add_heading(document, text, level):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    set_font(paragraph.add_run(text), size={1: 16, 2: 14, 3: 12}[level], color=HEADING)
    return paragraph


def add_activity(document, activity):
    add_heading(document, f"{activity['id']} {activity['title']}", 2)
    add_label(document, "Descrição Inicial:")
    document.add_paragraph(activity["description"])

    add_label(document, "Funcionalidades")
    for item in activity["functionalities"]:
        add_bullet(document, item)

    add_label(document, "Desafios")
    for item in activity["challenges"]:
        add_bullet(document, item)

    if activity.get("architecture"):
        add_label(document, "Arquitetura")
        for name, description in activity["architecture"]:
            add_component(document, name, description)

    if activity.get("flow"):
        add_label(document, "Fluxo esperado:")
        for item in activity["flow"]:
            add_number(document, item)

    add_label(document, "Metodologia")
    for stage, details in activity["methodology"]:
        add_heading(document, stage, 3)
        for item in details:
            add_bullet(document, item)

    add_label(document, "Entregáveis")
    for item in activity["deliverables"]:
        add_bullet(document, item)

    add_label(document, "Estimativa", activity["estimate"])

    add_label(document, "Dependências / Requisitos Básicos")
    for item in activity["dependencies"]:
        add_bullet(document, item)

    if activity.get("questions"):
        add_label(document, "Perguntas / Dúvidas")
        for item in activity["questions"]:
            add_bullet(document, item)


doc = Document()
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.49)
section.footer_distance = Inches(0.49)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for name, size, before, after in (
    ("Heading 1", 16, 14, 8),
    ("Heading 2", 14, 14, 8),
    ("Heading 3", 12, 10, 5),
):
    style = doc.styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = False
    style.font.color.rgb = RGBColor.from_string(HEADING)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.keep_with_next = True

for name in ("Plano Bullet", "Plano Number"):
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = normal
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

BULLET_ID, NUMBER_ID = create_numbering(doc)

# Abertura no mesmo padrão do PDF de referência
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(18)
title.paragraph_format.space_after = Pt(12)
title.paragraph_format.keep_with_next = True
set_font(title.add_run("Projeto Usinas de Geração de Energia"), size=20)

add_label(doc, "Descrição Inicial:")
doc.add_paragraph(
    "Este plano tem como objetivo definir as atividades para conceber, desenvolver, testar e transferir uma solução de segurança criptográfica aplicável a usinas de geração de energia. O projeto considera ambientes de tecnologia da informação e tecnologia operacional, com foco em disponibilidade, rastreabilidade, continuidade e evolução dos mecanismos criptográficos."
)
doc.add_paragraph(
    "O trabalho é organizado em duas frentes complementares. A primeira trata da modernização criptográfica, incluindo correções, proxy/gateway, inventário, criptoagilidade e desacoplamento das aplicações. A segunda trata da avaliação e integração de distribuição quântica de chaves (QKD) em laboratório e piloto controlado."
)

add_label(doc, "Fases:")
for item in (
    "2. Conceitual",
    "3. Desenvolvimento",
    "4. Testes",
    "5. Transferência de Tecnologia",
):
    add_bullet(doc, item)

add_label(doc, "Estimativa geral:", "42 a 56 semanas para concepção, POC e piloto. Uma implantação física de QKD pode acrescentar de 8 a 16 semanas, conforme aquisição, fibra, logística e homologação.")

# 2. Conceitual
add_heading(doc, "2. Conceitual", 1)
doc.add_paragraph(
    "A fase Conceitual define arquitetura, requisitos, governança, inventário, critérios de risco e estratégia de migração antes da implementação. As atividades podem ocorrer em paralelo e possuem duração de calendário estimada entre 10 e 14 semanas."
)

conceptual = [
    {
        "id": "2.1",
        "title": "Definição e Detalhamento da Arquitetura",
        "description": "Definir a arquitetura atual, a arquitetura-alvo e a arquitetura de transição para os ambientes de TI, TO e interconexões das usinas, identificando fluxos críticos, limites de confiança e pontos de aplicação da criptografia.",
        "functionalities": [
            "Mapear zonas, redes, centros de operação, datacenters, acesso remoto e enlaces entre unidades.",
            "Identificar aplicações, equipamentos, protocolos, certificados, chaves e serviços de segurança envolvidos nos fluxos.",
            "Definir componentes de proxy/gateway, abstração criptográfica, PKI, KMS/HSM, inventário e monitoração.",
            "Definir requisitos de disponibilidade, latência, capacidade, auditoria, recuperação e reversão.",
            "Produzir uma arquitetura de transição que permita coexistência entre mecanismos atuais e novos.",
        ],
        "challenges": [
            "Documentação incompleta ou desatualizada dos ambientes.",
            "Restrições de fabricantes e equipamentos legados.",
            "Necessidade de preservar funções de proteção e controle da planta.",
            "Dependências entre redes de TI, TO, telecomunicações e fornecedores externos.",
        ],
        "architecture": [
            ("Arquitetura Atual", "Representa ativos, fluxos, protocolos e controles criptográficos existentes."),
            ("Arquitetura-Alvo", "Define os componentes e interfaces que permitirão modernização, criptoagilidade e integração QKD."),
            ("Arquitetura de Transição", "Organiza coexistência, migração gradual, rollback e retirada de componentes legados."),
        ],
        "flow": [
            "Coletar diagramas, inventários e informações dos responsáveis.",
            "Modelar a situação atual e validar os fluxos críticos.",
            "Definir a arquitetura-alvo e os requisitos não funcionais.",
            "Definir etapas de transição, coexistência e reversão.",
            "Revisar e aprovar a arquitetura com segurança, engenharia e operação.",
        ],
        "methodology": [
            ("1. Levantamento", ["Entrevistas e workshops.", "Análise de diagramas, aplicações, redes e ativos."]),
            ("2. Modelagem", ["Diagramas de contexto, componentes e implantação.", "Catálogo de fluxos, protocolos e interfaces."]),
            ("3. Validação", ["Revisões multidisciplinares.", "Registro de decisões, riscos e pendências."]),
        ],
        "deliverables": ["Documento de arquitetura atual, alvo e de transição.", "Diagramas e catálogo de interfaces.", "Requisitos funcionais e não funcionais.", "Registro de decisões e riscos."],
        "estimate": "3 a 5 semanas.",
        "dependencies": ["Diagramas e inventários existentes.", "Disponibilidade dos especialistas de TI, TO e telecomunicações.", "Definição preliminar das usinas e aplicações do piloto."],
        "questions": ["Quais sistemas e enlaces fazem parte da primeira onda?", "Quais limites de latência e indisponibilidade são aceitáveis?", "Quais componentes não podem ser alterados diretamente?"],
    },
    {
        "id": "2.2",
        "title": "Governança e Criptoagilidade",
        "description": "Definir como a organização governará e executará substituições de algoritmos, bibliotecas, parâmetros, certificados e provedores criptográficos, preservando segurança, interoperabilidade e continuidade. Esta atividade é baseada no plano estratégico de criptoagilidade proposto pelo NIST CSWP 39-upd1.",
        "functionalities": [
            "Designar responsável e definir papéis para política, implantação, exceções e monitoramento.",
            "Integrar criptoagilidade à gestão de riscos e aos processos de mudança, aquisição e modernização.",
            "Definir política criptográfica com mecanismos permitidos, transitórios e proibidos.",
            "Transformar regras da política em perfis de configuração aplicáveis por automação, quando possível.",
            "Definir indicadores de cobertura, conformidade, tempo, custo e facilidade de migração.",
            "Definir níveis de maturidade e uma meta de evolução para as equipes e ambientes.",
        ],
        "challenges": [
            "Responsabilidades distribuídas entre segurança, aplicações, infraestrutura, operação e fornecedores.",
            "Políticas que não podem ser aplicadas automaticamente em sistemas legados.",
            "Necessidade de sincronizar a retirada de algoritmos entre todas as partes.",
            "Risco de aumentar a complexidade e a superfície de ataque com opções excessivas.",
        ],
        "architecture": [
            ("Governança", "Define responsáveis, decisões, comunicação, exceções e acompanhamento executivo."),
            ("Política Criptográfica", "Define padrões, algoritmos, protocolos, parâmetros e prazos permitidos."),
            ("Perfis Técnicos", "Traduzem a política para configurações específicas de aplicações, gateways e ferramentas."),
            ("Indicadores", "Medem a postura atual e a capacidade real de executar transições."),
        ],
        "flow": [
            "Identificar requisitos, partes interessadas e responsável pelo programa.",
            "Definir política e perfis técnicos iniciais.",
            "Relacionar a política aos ativos e riscos do inventário.",
            "Definir indicadores, metas e processo de exceção.",
            "Aprovar e comunicar a política às equipes e fornecedores.",
        ],
        "methodology": [
            ("1. Governança", ["Definir RACI e fóruns de decisão.", "Integrar criptoagilidade ao registro de riscos."]),
            ("2. Política", ["Definir mecanismos permitidos e proibidos.", "Criar perfis técnicos e regras de transição."]),
            ("3. Maturidade", ["Avaliar situação atual.", "Definir indicadores e meta de evolução."]),
        ],
        "deliverables": ["Política criptográfica.", "RACI e fluxo de decisão.", "Perfis técnicos iniciais.", "Indicadores e avaliação de maturidade.", "Processo de exceção e comunicação."],
        "estimate": "3 a 5 semanas.",
        "dependencies": ["Patrocínio executivo e técnico.", "Participação de segurança, arquitetura, operação, compras e jurídico.", "Requisitos normativos e internos aplicáveis."],
        "questions": ["Quem aprova a retirada de um algoritmo?", "Como a política será transformada em configurações técnicas?", "Qual nível de maturidade deve ser atingido no piloto?"],
    },
    {
        "id": "2.3",
        "title": "Inventário Lógico, Físico e Criptográfico",
        "description": "Construir uma linha de base dos dados, ativos, aplicações, equipamentos, protocolos, bibliotecas, certificados e chaves que utilizam criptografia, adotando uma visão centrada na criticidade dos dados e dos casos de uso.",
        "functionalities": [
            "Relacionar dados em repouso, em trânsito e em uso aos controles criptográficos aplicados.",
            "Identificar algoritmos, protocolos, bibliotecas, versões, comprimentos de chave e certificados.",
            "Registrar PKI, autoridades certificadoras, KMS, HSM, cofres e locais de armazenamento de chaves.",
            "Relacionar ativos aos proprietários, fabricantes, criticidade, exposição e dificuldade de migração.",
            "Identificar lacunas, itens desconhecidos, fim de suporte e dependências externas.",
        ],
        "challenges": ["Ativos não cadastrados ou sem proprietário.", "Protocolos proprietários e criptografia incorporada em firmware.", "Limitações de varredura ativa em ambientes de TO.", "Duplicidade e inconsistência entre diferentes inventários."],
        "architecture": [
            ("Fontes de Inventário", "CMDB, gestão de ativos, redes, repositórios, pipelines, PKI, KMS/HSM e fabricantes."),
            ("Modelo de Dados", "Relaciona dado, ativo, aplicação, protocolo, algoritmo, chave, certificado, proprietário e risco."),
            ("Visão de Risco", "Prioriza itens por criticidade, exposição, obsolescência e capacidade de migração."),
        ],
        "flow": ["Definir o modelo de dados.", "Coletar informações de fontes existentes.", "Executar descoberta passiva ou homologada.", "Validar os resultados com os proprietários.", "Classificar riscos e lacunas."],
        "methodology": [
            ("1. Inventário Inicial", ["Consolidação de fontes existentes.", "Entrevistas e validação com proprietários."]),
            ("2. Descoberta", ["Análise de certificados, endpoints, código, bibliotecas e equipamentos.", "Coleta passiva em TO sempre que possível."]),
            ("3. Classificação", ["Criticidade dos dados e ativos.", "Risco, exposição e dificuldade de migração."]),
        ],
        "deliverables": ["Modelo de dados do inventário.", "Inventário criptográfico inicial.", "Mapa de risco e lacunas.", "Plano de atualização contínua."],
        "estimate": "4 a 7 semanas.",
        "dependencies": ["Acesso de leitura às fontes de ativos.", "Aprovação das técnicas de descoberta em TO.", "Participação de proprietários e fornecedores."],
        "questions": ["Qual percentual dos ativos prioritários deve ser coberto?", "Quais métodos de descoberta são permitidos em TO?", "Onde o inventário oficial será mantido?"],
    },
    {
        "id": "2.4",
        "title": "Levantamento de Ferramentas e Automação",
        "description": "Avaliar se as ferramentas existentes conseguem descobrir, caracterizar, aplicar políticas e monitorar criptografia, definindo automações para reduzir levantamentos e verificações manuais.",
        "functionalities": [
            "Avaliar CMDB, gestão de ativos, vulnerabilidades, configurações, logs, repositórios e pipelines.",
            "Verificar descoberta de algoritmos, protocolos, bibliotecas, certificados e comprimentos de chave.",
            "Definir integrações e modelo comum de dados.",
            "Definir validações de política em pipelines e ambientes.",
            "Definir alertas de expiração, mecanismo proibido, biblioteca vulnerável e ausência de proprietário.",
        ],
        "challenges": ["Ferramentas isoladas e modelos de dados incompatíveis.", "Cobertura limitada de equipamentos e firmware.", "Falsos positivos e duplicidades.", "Restrições de credenciais e coleta em redes segregadas."],
        "architecture": [
            ("Coletores", "Obtêm evidências de fontes autorizadas."),
            ("Normalização", "Relaciona evidências ao modelo do inventário."),
            ("Motor de Política", "Compara o estado encontrado com os perfis permitidos."),
            ("Painéis e Alertas", "Apresentam risco, conformidade, expiração e evolução da migração."),
        ],
        "flow": ["Mapear ferramentas e capacidades.", "Comparar cobertura com o inventário desejado.", "Priorizar lacunas.", "Definir integrações e automações.", "Validar uma coleta automatizada em ambiente controlado."],
        "methodology": [
            ("1. Avaliação", ["Matriz ferramenta x capacidade.", "Análise de cobertura e qualidade dos dados."]),
            ("2. Desenho", ["Arquitetura de integração.", "Regras de política e alertas."]),
            ("3. Piloto", ["Automatizar uma fonte prioritária.", "Medir cobertura e falsos positivos."]),
        ],
        "deliverables": ["Matriz de capacidades e lacunas.", "Arquitetura de integração.", "Backlog de automação.", "Regras iniciais de política e alertas."],
        "estimate": "2 a 4 semanas.",
        "dependencies": ["Acesso às ferramentas corporativas.", "Inventário inicial.", "Disponibilidade das equipes responsáveis."],
        "questions": ["Quais ferramentas podem ser reutilizadas?", "Quais lacunas exigem aquisição?", "Como as regras serão aprovadas e atualizadas?"],
    },
    {
        "id": "2.5",
        "title": "Planejamento Criptográfico e Roadmap",
        "description": "Transformar inventário, política e riscos em uma sequência priorizada de correções e migrações, considerando a vida útil dos dados, criticidade operacional, capacidade dos ativos e ciclos de modernização.",
        "functionalities": ["Priorizar ativos e casos de uso.", "Definir ondas de migração e critérios de entrada e saída.", "Definir coexistência, controle compensatório e rollback.", "Definir indicadores, responsáveis, custos e marcos.", "Planejar revisões periódicas do inventário, da política e dos riscos."],
        "challenges": ["Dependência de fornecedores e janelas operacionais.", "Ativos que não podem ser atualizados.", "Conflito entre urgência de segurança e risco de indisponibilidade.", "Migrações que exigem interoperabilidade temporária com legado."],
        "architecture": [
            ("Onda 1", "Correções prioritárias e POCs em ambientes controlados."),
            ("Onda 2", "Pilotos integrados e aplicações de maior valor."),
            ("Onda 3", "Escalonamento, retirada de legado e melhoria contínua."),
        ],
        "flow": ["Aplicar a política ao inventário.", "Calcular risco e capacidade de migração.", "Definir ações e controles compensatórios.", "Organizar ondas e marcos.", "Aprovar orçamento, responsáveis e critérios de aceite."],
        "methodology": [
            ("1. Priorização", ["Criticidade, exposição e vida útil do dado.", "Tempo, custo e facilidade de migração."]),
            ("2. Planejamento", ["Ondas, pilotos e dependências.", "Rollback e controles compensatórios."]),
            ("3. Governança", ["Indicadores e marcos.", "Revisão contínua de riscos e políticas."]),
        ],
        "deliverables": ["Roadmap de 12, 24 e 36 meses.", "Matriz de priorização.", "Plano de ondas e pilotos.", "Indicadores e critérios de aceite."],
        "estimate": "3 a 5 semanas.",
        "dependencies": ["Inventário e política criptográfica.", "Disponibilidade dos donos dos ativos.", "Informações de ciclo de vida e orçamento."],
        "questions": ["Quais ativos devem ser corrigidos primeiro?", "Quais precisam de substituição?", "Quais controles compensatórios são aceitáveis?"],
    },
    {
        "id": "2.6",
        "title": "Cadeia de Suprimentos e Aquisições",
        "description": "Incorporar requisitos de criptoagilidade às relações com fabricantes, integradores e fornecedores, identificando componentes que podem ser atualizados e aqueles que exigirão substituição ou mitigação.",
        "functionalities": ["Mapear dependências de hardware, firmware, software, bibliotecas, protocolos e serviços.", "Solicitar a composição criptográfica e a capacidade de atualização dos produtos.", "Definir requisitos de aquisição e homologação.", "Classificar fim de suporte, substituição, mitigação e modernização.", "Definir responsabilidades e prazos de correção contratual."],
        "challenges": ["Informações técnicas incompletas dos fornecedores.", "Produtos sem mecanismo de atualização modular.", "Dependência de padrões e validações ainda em evolução.", "Ciclos longos de aquisição e substituição industrial."],
        "architecture": [
            ("Fabricantes", "Informam composição, suporte, atualização e limitações criptográficas."),
            ("Aquisição", "Inclui requisitos de atualização, inventário, suporte e migração."),
            ("Engenharia e Segurança", "Avaliam compatibilidade, risco e homologação."),
        ],
        "flow": ["Identificar fornecedores críticos.", "Aplicar questionário técnico.", "Avaliar capacidade de atualização.", "Definir requisitos contratuais e de homologação.", "Registrar riscos e planos de tratamento."],
        "methodology": [
            ("1. Levantamento", ["Produtos, versões, contratos e fim de suporte.", "Capacidade de atualização e exportação de inventário."]),
            ("2. Requisitos", ["Cláusulas de criptoagilidade.", "Critérios de POC, homologação e correção."]),
            ("3. Tratamento", ["Atualizar, substituir ou compensar.", "Registrar responsáveis e prazos."]),
        ],
        "deliverables": ["Matriz fornecedor x componente.", "Questionário técnico.", "Requisitos de aquisição e homologação.", "Plano para itens em fim de vida ou sem atualização."],
        "estimate": "3 a 6 semanas.",
        "dependencies": ["Lista de fornecedores e contratos.", "Participação de compras, jurídico, engenharia e segurança.", "Acesso à documentação dos produtos."],
        "questions": ["Quais produtos suportam novos algoritmos sem substituição?", "Quais evidências devem ser exigidas?", "Qual prazo de correção será contratualmente aceitável?"],
    },
    {
        "id": "2.7",
        "title": "Concepção da Arquitetura QKD",
        "description": "Avaliar onde a distribuição quântica de chaves agrega valor e definir uma arquitetura de laboratório ou piloto integrada à gestão de chaves existente, mantendo autenticação, proteção dos endpoints e operação híbrida.",
        "functionalities": ["Selecionar enlaces candidatos por criticidade, distância, fibra e atenuação.", "Definir nós QKD, canal quântico, canal clássico autenticado e pontos de confiança.", "Definir integração com KMS/HSM, encryptor, VPN, TLS ou aplicação.", "Definir consumo, estoque, rotação, expiração e descarte de chaves.", "Definir fallback e operação híbrida quando o enlace QKD estiver indisponível."],
        "challenges": ["Disponibilidade e qualidade da fibra.", "Aquisição e logística de equipamentos.", "Integração entre interfaces de fornecedores.", "Manutenção da continuidade quando não houver chaves QKD disponíveis."],
        "architecture": [
            ("Nós QKD", "Geram material de chave correlacionado nos extremos do enlace."),
            ("Canal Clássico Autenticado", "Suporta os protocolos auxiliares e precisa de autenticação forte."),
            ("Gerenciador de Chaves", "Recebe, armazena, controla e entrega chaves para o caso de uso."),
            ("Aplicação ou Encryptor", "Utiliza as chaves para proteger o tráfego ou dados selecionados."),
        ],
        "flow": ["Selecionar enlace e caso de uso.", "Definir interfaces e requisitos.", "Montar simulador ou laboratório.", "Integrar gestão e consumo de chaves.", "Validar operação nominal, fallback e indicadores."],
        "methodology": [
            ("1. Viabilidade", ["Distância, fibra, atenuação e criticidade.", "Interfaces e equipamentos disponíveis."]),
            ("2. Arquitetura", ["Nós, canais, KMS/HSM e consumidor.", "Política de chaves, autenticação e fallback."]),
            ("3. Plano de POC", ["Ambiente, indicadores e testes.", "Critérios para decidir um piloto físico."]),
        ],
        "deliverables": ["Arquitetura QKD conceitual.", "Matriz de enlaces candidatos.", "Requisitos de integração e segurança.", "Plano e critérios da POC QKD."],
        "estimate": "3 a 5 semanas.",
        "dependencies": ["Dados dos enlaces.", "Participação de telecomunicações e redes.", "Acesso a simulador, laboratório ou fornecedor QKD."],
        "questions": ["Qual enlace possui melhor relação entre valor e viabilidade?", "Qual aplicação consumirá as chaves?", "Qual será o mecanismo de fallback?"],
    },
]

for item in conceptual:
    add_activity(doc, item)

# 3. Desenvolvimento
add_heading(doc, "3. Desenvolvimento", 1)
doc.add_paragraph(
    "A fase de Desenvolvimento implementa as correções, componentes e integrações aprovadas. O trabalho começa em laboratório e ambiente de integração, com automação, versionamento, observabilidade e rollback. A duração estimada é de 16 a 24 semanas, com frentes paralelas."
)

development = [
    {
        "id": "3.1",
        "title": "Correções e Modernização Criptográfica",
        "description": "Corrigir configurações inseguras, dependências obsoletas e falhas de gestão de certificados e chaves identificadas no inventário.",
        "functionalities": ["Atualizar bibliotecas e provedores.", "Remover protocolos, algoritmos e parâmetros proibidos.", "Corrigir validação e ciclo de vida de certificados.", "Retirar segredos fixos de código e configuração.", "Adicionar testes de regressão e evidências antes/depois."],
        "challenges": ["Compatibilidade com aplicações e fabricantes.", "Risco de regressão.", "Janelas de mudança restritas.", "Itens que exigem substituição e não apenas atualização."],
        "architecture": [("Aplicação", "Recebe a correção ou nova configuração."), ("PKI/KMS/HSM", "Fornece certificados, chaves e proteção adequada."), ("Pipeline", "Valida dependências e política antes da implantação.")],
        "flow": ["Selecionar a onda.", "Reproduzir e registrar o estado atual.", "Implementar a correção.", "Executar regressão e segurança.", "Implantar em piloto com rollback disponível."],
        "methodology": [("1. Correção", ["Atualização, configuração e integração segura."]), ("2. Validação", ["Regressão, segurança e desempenho."]), ("3. Implantação", ["Mudança controlada, evidências e rollback."])],
        "deliverables": ["Pacotes de correção.", "Testes automatizados.", "Evidências antes/depois.", "Registro de exceções."],
        "estimate": "4 a 8 semanas por onda de 3 a 5 aplicações.",
        "dependencies": ["Aplicações e ambientes de teste.", "Proprietários e fornecedores.", "Política e prioridades aprovadas."],
        "questions": ["Quais correções podem ser feitas por configuração?", "Quais exigem alteração de código?", "Quais dependem do fabricante?"],
    },
    {
        "id": "3.2",
        "title": "Proxy/Gateway Criptográfico",
        "description": "Implementar um proxy ou gateway para centralizar políticas, terminação ou encaminhamento de sessões, certificados, rotação de chaves e métricas, principalmente para sistemas legados que não podem ser alterados diretamente.",
        "functionalities": ["Implementar comunicação cliente-proxy-servidor.", "Definir terminação, recriptografia ou encaminhamento por fluxo.", "Aplicar TLS, autenticação mútua e cadeia de confiança.", "Aplicar políticas e rotação de modo centralizado.", "Coletar handshake, latência, algoritmo, erros e expiração.", "Implementar alta disponibilidade, health checks e fallback."],
        "challenges": ["Compatibilidade com protocolos e aplicações.", "Risco de ponto único de falha.", "Gestão de certificados e chaves.", "Capacidade, latência e tratamento de falhas."],
        "architecture": [("Cliente", "Inicia a comunicação e valida o certificado apresentado."), ("Proxy/Gateway", "Aplica política, protege o fluxo, encaminha o tráfego e coleta métricas."), ("PKI/KMS/HSM", "Emite certificados e protege chaves."), ("Servidor", "Recebe o fluxo protegido ou reencaminhado.")],
        "flow": ["Provisionar certificados e políticas.", "Cliente conecta ao proxy.", "Proxy negocia a sessão segura.", "Proxy conecta ao servidor.", "Tráfego é encaminhado e métricas são registradas.", "Falhas acionam alta disponibilidade ou rollback."],
        "methodology": [("1. Proxy Básico", ["Comunicação convencional e baseline."]), ("2. Integração", ["PKI/KMS, política, telemetria e alta disponibilidade."]), ("3. Agilidade", ["Troca de perfil, coexistência e retirada de legado."])],
        "deliverables": ["Proxy/gateway implantável.", "Configurações versionadas.", "Integração com PKI/KMS e monitoração.", "Runbook e rollback."],
        "estimate": "6 a 10 semanas.",
        "dependencies": ["Fluxos do piloto.", "Ambiente cliente-proxy-servidor.", "Certificados e serviços de chaves de teste."],
        "questions": ["O proxy terminará ou apenas encaminhará o TLS?", "Qual topologia de alta disponibilidade será usada?", "Quais protocolos não podem passar pelo gateway?"],
    },
    {
        "id": "3.3",
        "title": "Desacoplamento das Aplicações e Crypto API",
        "description": "Separar regras de negócio das implementações criptográficas por meio de adaptadores, biblioteca corporativa, API ou serviço, permitindo substituir provedores e algoritmos com menor alteração nas aplicações.",
        "functionalities": ["Mapear chamadas criptográficas e dependências.", "Definir interface simples e padrões seguros.", "Implementar provedores intercambiáveis.", "Versionar envelopes e metadados de algoritmo.", "Permitir coexistência de leitura e transição gradual.", "Aplicar política por configuração autorizada e auditável."],
        "challenges": ["Formatos de dados e assinaturas existentes.", "APIs excessivamente genéricas ou difíceis de usar.", "Compatibilidade durante migração.", "Dependências de hardware, kernel ou HSM."],
        "architecture": [("Aplicação", "Solicita operações de alto nível sem incorporar detalhes desnecessários."), ("Crypto API", "Oferece contratos estáveis e valida a política."), ("Provedor", "Implementa algoritmos em software, hardware ou serviço."), ("Configuração de Política", "Seleciona opções autorizadas e prazos de transição.")],
        "flow": ["Mapear o acoplamento atual.", "Definir o contrato da API.", "Implementar provedor atual e alternativo.", "Adaptar uma aplicação piloto.", "Trocar o provedor em laboratório.", "Validar coexistência e rollback."],
        "methodology": [("1. Plano de Mudança", ["Módulos, releases, compatibilidade e rollback."]), ("2. Implementação", ["API, provedores e configuração de política."]), ("3. Demonstração", ["Troca de provedor sem alteração da regra de negócio."])],
        "deliverables": ["Plano de desacoplamento.", "Crypto API/SDK ou serviço.", "Dois provedores de demonstração.", "Adaptador da aplicação piloto.", "Testes de contrato e rollback."],
        "estimate": "6 a 10 semanas para a plataforma e 3 a 6 semanas por aplicação piloto.",
        "dependencies": ["Código-fonte e especialistas da aplicação.", "Política e contratos aprovados.", "Provedores e ambientes de integração."],
        "questions": ["A melhor estratégia é API, biblioteca, serviço ou proxy?", "Como os dados antigos continuarão legíveis?", "Quais escolhas permanecerão visíveis ao desenvolvedor?"],
    },
    {
        "id": "3.4",
        "title": "Automação do Inventário e da Política",
        "description": "Implementar coleta recorrente, normalização, avaliação de conformidade e alertas para acompanhar continuamente o estado da criptografia e a evolução das migrações.",
        "functionalities": ["Criar conectores para fontes autorizadas.", "Normalizar e relacionar evidências aos ativos.", "Aplicar perfis de política consumíveis por máquina.", "Gerar alertas e recomendações.", "Publicar painéis de risco, cobertura, conformidade e expiração."],
        "challenges": ["Qualidade e duplicidade dos dados.", "Credenciais e redes segregadas.", "Falsos positivos.", "Cobertura limitada de firmware e sistemas proprietários."],
        "architecture": [("Coletores", "Obtêm evidências sem afetar a operação."), ("Base Normalizada", "Mantém relações e histórico."), ("Motor de Avaliação", "Compara evidências com a política."), ("Painel", "Apoia operação, gestão e auditoria.")],
        "flow": ["Coletar dados.", "Normalizar e correlacionar.", "Aplicar política.", "Classificar risco e prioridade.", "Alertar e abrir ações.", "Repetir e medir evolução."],
        "methodology": [("1. Conectores", ["Priorizar fontes de maior valor."]), ("2. Política", ["Implementar regras e exceções."]), ("3. Operação", ["Painéis, alertas e tratamento de erros."])],
        "deliverables": ["Coletores automatizados.", "Base consolidada.", "Regras de avaliação.", "Painéis e alertas.", "Procedimento de operação."],
        "estimate": "4 a 7 semanas para o piloto.",
        "dependencies": ["Modelo do inventário.", "Perfis de política.", "Acessos de leitura e integração."],
        "questions": ["Quais fontes serão automatizadas primeiro?", "Qual taxa de falsos positivos é aceitável?", "Quem tratará cada alerta?"],
    },
    {
        "id": "3.5",
        "title": "Desenvolvimento QKD",
        "description": "Montar o ambiente QKD, integrar a entrega de chaves ao KMS ou gerenciador, conectar o caso de uso e implementar operação híbrida, fallback e observabilidade.",
        "functionalities": ["Montar nós e canais em simulador ou laboratório.", "Implementar adaptador da interface QKD para gestão de chaves.", "Integrar encryptor, VPN, TLS ou aplicação selecionada.", "Implementar consumo, estoque, rotação e descarte.", "Implementar fallback, failback e alarmes.", "Coletar disponibilidade, taxa, estoque, consumo e erros sem registrar material de chave."],
        "challenges": ["Interoperabilidade de interfaces.", "Sincronização e disponibilidade de chaves.", "Integração com KMS/HSM e aplicação.", "Falha do enlace e retorno controlado.", "Aquisição e implantação física."],
        "architecture": [("Sistema QKD", "Gera e sincroniza material de chave."), ("Adaptador QKD", "Autentica, controla e traduz a interface."), ("KMS/Gerenciador", "Armazena e entrega chaves de acordo com a política."), ("Consumidor", "Utiliza a chave no mecanismo de proteção selecionado."), ("Observabilidade", "Registra estado, métricas e falhas sem expor chaves.")],
        "flow": ["Gerar e disponibilizar chaves.", "Adaptador consulta ou recebe a chave.", "KMS controla o ciclo de vida.", "Consumidor aplica a chave.", "Métricas são registradas.", "Indisponibilidade aciona o fallback aprovado."],
        "methodology": [("1. Ambiente", ["Simulador ou laboratório e conectividade."]), ("2. Integração", ["Adaptador, KMS e consumidor."]), ("3. Continuidade", ["Fallback, failback, alarmes e recuperação."]), ("4. Piloto", ["Métricas, evidências e decisão sobre implantação física."])],
        "deliverables": ["Ambiente QKD reproduzível.", "Adaptador de chaves.", "Integração ponta a ponta.", "Operação híbrida e fallback.", "Painel e runbook."],
        "estimate": "12 a 20 semanas para integração em laboratório. Implantação física pode acrescentar 8 a 16 semanas.",
        "dependencies": ["Simulador ou equipamento QKD.", "KMS/HSM ou gerenciador de testes.", "Caso de uso e enlace selecionados.", "Fornecedor ou laboratório especializado."],
        "questions": ["Qual interface de chaves será usada?", "Qual taxa de consumo o caso exige?", "Como a aplicação se comportará durante o fallback?"],
    },
]

for item in development:
    add_activity(doc, item)

# 4. Testes
add_heading(doc, "4. Testes", 1)
doc.add_paragraph(
    "A fase de Testes valida requisitos, segurança, interoperabilidade, desempenho, continuidade e capacidade real de troca. Os testes começam por componente e avançam para integração e piloto. A duração estimada é de 8 a 14 semanas."
)

tests = [
    {
        "id": "4.1",
        "title": "Testes Funcionais, Integração e Interoperabilidade",
        "description": "Validar fluxos ponta a ponta, certificados, chaves, políticas, proxy, Crypto API, inventário, interfaces QKD e coexistência com componentes legados.",
        "functionalities": ["Validar cenários nominais e negativos.", "Testar versões, clientes, servidores, bibliotecas e equipamentos.", "Validar cadeia de confiança, rotação, revogação e expiração.", "Validar negociação e conjunto mínimo de interoperabilidade.", "Verificar que opções proibidas não são selecionadas."],
        "challenges": ["Matriz extensa de versões e fornecedores.", "Diferenças de implementação.", "Ambientes de teste pouco representativos.", "Necessidade de reproduzir falhas de modo seguro."],
        "flow": ["Preparar ambiente e pré-condições.", "Executar fluxo nominal.", "Executar incompatibilidades e falhas.", "Coletar evidências.", "Corrigir e repetir.", "Comparar com critérios de aceite."],
        "methodology": [("1. Casos de Teste", ["Requisitos, riscos e rastreabilidade."]), ("2. Execução", ["Automação, evidências e defeitos."]), ("3. Regressão", ["Repetição após correções."])],
        "deliverables": ["Plano e casos de teste.", "Evidências.", "Relatório de defeitos.", "Matriz de interoperabilidade."],
        "estimate": "3 a 5 semanas.",
        "dependencies": ["Componentes integrados.", "Ambiente estável.", "Massa e certificados de teste."],
        "questions": ["Qual conjunto mínimo precisa interoperar?", "Quais versões legadas serão mantidas temporariamente?"],
    },
    {
        "id": "4.2",
        "title": "Testes de Criptoagilidade e Transição",
        "description": "Demonstrar que a organização consegue trocar algoritmos, provedores, certificados ou perfis dentro de prazo conhecido, sem alteração indevida da aplicação e com rollback auditável.",
        "functionalities": ["Trocar provedor ou perfil criptográfico.", "Introduzir uma opção nova e retirar uma antiga.", "Testar coexistência e leitura de dados existentes.", "Testar proteção contra downgrade.", "Medir tempo, esforço, custo, indisponibilidade e quantidade de alterações.", "Validar perfis de política e atualização automatizada."],
        "challenges": ["Dependências ocultas.", "Formatos de dados sem versionamento.", "Negociação insegura ou fallback silencioso.", "Rollback não exercitado."],
        "flow": ["Registrar o estado inicial.", "Publicar nova política ou provedor.", "Executar a transição.", "Verificar aplicações e dados.", "Retirar a opção antiga.", "Executar rollback e calcular indicadores."],
        "methodology": [("1. Exercício de Transição", ["Cenário planejado e janela controlada."]), ("2. Exercício de Urgência", ["Simular retirada acelerada por vulnerabilidade."]), ("3. Medição", ["Tempo, custo, impacto e evidências."])],
        "deliverables": ["Roteiro de transição.", "Evidências de troca e rollback.", "Indicadores de agilidade.", "Lições e backlog de melhoria."],
        "estimate": "3 a 5 semanas.",
        "dependencies": ["Dois perfis ou provedores disponíveis.", "Política e critérios de aceite.", "Telemetria e ambiente controlado."],
        "questions": ["Quanto tempo a transição deve levar?", "Qual indisponibilidade é aceitável?", "Quais partes ainda exigem intervenção manual?"],
    },
    {
        "id": "4.3",
        "title": "Desempenho, Resiliência e Segurança",
        "description": "Avaliar o impacto das mudanças e comprovar continuidade, capacidade e proteção contra falhas e ataques relevantes.",
        "functionalities": ["Medir handshake, latência, vazão, CPU, memória e tamanho de mensagens/certificados.", "Testar carga, picos e filas.", "Simular falha de proxy, PKI, KMS/HSM, rede e fornecedor.", "Executar failover, failback, recuperação e rollback.", "Revisar código, configuração, acessos, segredos, logs e negociação contra downgrade."],
        "challenges": ["Baselines incompletas.", "Diferenças entre laboratório e planta.", "Falhas que não podem ser simuladas em produção.", "Algoritmos e certificados maiores exigindo mais recursos."],
        "flow": ["Registrar baseline.", "Aplicar carga representativa.", "Introduzir falhas controladas.", "Medir recuperação.", "Executar revisão de segurança.", "Comparar com metas e riscos."],
        "methodology": [("1. Desempenho", ["Baseline e cenários de carga."]), ("2. Continuidade", ["Falhas, recuperação e rollback."]), ("3. Segurança", ["Código, configuração, acesso e superfícies expostas."])],
        "deliverables": ["Relatório de desempenho.", "Relatório de continuidade.", "Relatório de segurança e hardening.", "Riscos residuais e recomendações."],
        "estimate": "4 a 6 semanas.",
        "dependencies": ["Metas de desempenho e recuperação.", "Ambiente representativo.", "Ferramentas e equipe de testes."],
        "questions": ["Qual degradação é aceitável?", "Qual RTO deve ser atingido?", "Quais falhas precisam de teste em laboratório dedicado?"],
    },
    {
        "id": "4.4",
        "title": "Testes QKD e Aceite do Piloto",
        "description": "Validar o ciclo de vida das chaves QKD, a disponibilidade do enlace, a integração com o consumidor, o fallback e o impacto operacional em ambiente representativo.",
        "functionalities": ["Medir taxa de geração, estoque, consumo e disponibilidade.", "Testar sincronização, expiração e descarte.", "Simular perda do enlace, canal clássico, nó ou estoque de chaves.", "Validar fallback e retorno controlado.", "Confirmar que logs não expõem chaves.", "Executar o piloto com acompanhamento de operação e segurança."],
        "challenges": ["Variação da qualidade do enlace.", "Estoque insuficiente para o consumo.", "Diferenças entre simulador e equipamento físico.", "Coordenação com janelas e equipes da usina."],
        "flow": ["Validar o enlace nominal.", "Consumir chaves no caso de uso.", "Medir indicadores.", "Introduzir falhas.", "Executar fallback e retorno.", "Formalizar aceite ou pendências."],
        "methodology": [("1. Validação Técnica", ["Taxa, estoque, consumo e falhas."]), ("2. Validação Operacional", ["Alertas, runbooks e escalonamento."]), ("3. Aceite", ["Indicadores, riscos e decisão sobre expansão."])],
        "deliverables": ["Evidências QKD.", "Relatório de disponibilidade e consumo.", "Evidências de fallback/failback.", "Termo de aceite ou plano de ajustes."],
        "estimate": "3 a 6 semanas.",
        "dependencies": ["Integração QKD concluída.", "Indicadores e limites aprovados.", "Participação das equipes de operação e fornecedor."],
        "questions": ["A taxa de chaves atende ao caso de uso?", "O fallback preserva a continuidade?", "O piloto justifica implantação física ou expansão?"],
    },
]

for item in tests:
    add_activity(doc, item)

# 5. Transferência
add_heading(doc, "5. Transferência de Tecnologia", 1)
doc.add_paragraph(
    "A transferência de tecnologia prepara as equipes para instalar, operar, monitorar, manter e evoluir a solução sem dependência excessiva da equipe de desenvolvimento ou de fornecedores. A duração estimada é de 6 a 10 semanas e pode se sobrepor aos testes."
)

transfer = [
    {
        "id": "5.1",
        "title": "Documentação e Capacitação",
        "description": "Consolidar arquitetura, decisões, configurações, procedimentos e material de treinamento para os diferentes perfis envolvidos.",
        "functionalities": ["Documentar arquitetura e interfaces.", "Documentar instalação, configuração, atualização e rollback.", "Criar runbooks de certificados, chaves, proxy, Crypto API, inventário e QKD.", "Preparar módulos para arquitetura, desenvolvimento, segurança, infraestrutura, operação e gestão.", "Registrar limites, riscos e dependências de fornecedores."],
        "challenges": ["Documentação que se torna obsoleta rapidamente.", "Públicos com diferentes níveis de conhecimento.", "Dependência de procedimentos tácitos.", "Disponibilidade das equipes para treinamento."],
        "flow": ["Consolidar documentos preliminares.", "Revisar com especialistas.", "Preparar trilhas por perfil.", "Executar aulas e demonstrações.", "Avaliar compreensão e atualizar o material."],
        "methodology": [("1. Documentação", ["Arquitetura, engenharia, operação e segurança."]), ("2. Capacitação", ["Conteúdo por perfil e estudos de caso."]), ("3. Avaliação", ["Exercícios, dúvidas e lacunas."])],
        "deliverables": ["Documentação final.", "Runbooks.", "Materiais de treinamento.", "Registro de participação e avaliação."],
        "estimate": "3 a 5 semanas.",
        "dependencies": ["Solução estável.", "Documentação técnica preliminar.", "Disponibilidade de instrutores e participantes."],
        "questions": ["Quais equipes precisam de cada trilha?", "Quais procedimentos exigem exercício obrigatório?"],
    },
    {
        "id": "5.2",
        "title": "Laboratórios, Handover e Operação Assistida",
        "description": "Transferir a execução prática das rotinas, validar a autonomia das equipes e acompanhar os primeiros ciclos operacionais controlados.",
        "functionalities": ["Executar instalação e recuperação em laboratório.", "Exercitar rotação, revogação, mudança de perfil e rollback.", "Exercitar falha de proxy, KMS/HSM e QKD.", "Transferir dashboards, acessos, SLAs, escalonamento e governança de mudança.", "Acompanhar a operação e tratar pendências."],
        "challenges": ["Falta de ambiente seguro para exercícios.", "Acessos e responsabilidades ainda não transferidos.", "Dependência de fornecedor em incidentes.", "Pouco tempo para repetir procedimentos críticos."],
        "flow": ["Preparar laboratórios.", "Executar exercícios guiados.", "Executar novamente sob responsabilidade da equipe de destino.", "Transferir acessos e rotinas.", "Iniciar operação assistida.", "Avaliar autonomia e formalizar aceite."],
        "methodology": [("1. Laboratório", ["Procedimentos nominais e de falha."]), ("2. Handover", ["Acessos, dashboards, SLAs e escalonamento."]), ("3. Operação Assistida", ["Acompanhamento, indicadores e correções."])],
        "deliverables": ["Laboratórios executados.", "Checklist de handover.", "Evidências de autonomia.", "Relatório de operação assistida.", "Aceite final e backlog de evolução."],
        "estimate": "4 a 8 semanas.",
        "dependencies": ["Runbooks aprovados.", "Ambiente de laboratório.", "Equipes e acessos definidos.", "Processo de suporte e escalonamento."],
        "questions": ["Quais tarefas precisam ser executadas sem apoio?", "Por quanto tempo haverá operação assistida?", "Quais pendências impedem o aceite final?"],
    },
]

for item in transfer:
    add_activity(doc, item)

add_heading(doc, "6. Cronograma Consolidado", 1)
add_label(doc, "Duração estimada por fase:")
for item in (
    "Conceitual: 10 a 14 semanas.",
    "Desenvolvimento: 16 a 24 semanas.",
    "Testes: 8 a 14 semanas.",
    "Transferência de Tecnologia: 6 a 10 semanas.",
    "Prazo total considerando sobreposição das fases: 42 a 56 semanas.",
    "Implantação física QKD: possível acréscimo de 8 a 16 semanas.",
):
    add_bullet(doc, item)

add_label(doc, "Premissas das estimativas:")
for item in (
    "Execução paralela por equipe multidisciplinar.",
    "Disponibilidade de ambiente de laboratório e integração.",
    "Acesso aos responsáveis, ativos, aplicações e fornecedores.",
    "Janelas controladas para mudanças e testes em ambientes representativos.",
    "Reestimativa ao final da fase Conceitual e após a primeira POC integrada.",
):
    add_bullet(doc, item)

add_heading(doc, "7. Referência Técnica", 1)
doc.add_paragraph(
    "National Institute of Standards and Technology (NIST). Considerations for Achieving Crypto Agility: Strategies and Practices. NIST CSWP 39-upd1, incluindo atualizações de 29 de junho de 2026. DOI: 10.6028/NIST.CSWP.39-upd1. Disponível em: https://nvlpubs.nist.gov/nistpubs/CSWP/NIST.CSWP.39-upd1.pdf."
)

doc.core_properties.title = "Projeto Usinas de Geração de Energia"
doc.core_properties.subject = "Plano de atividades em formato semelhante ao Plano 0"
doc.core_properties.author = "Equipe do Projeto"
doc.core_properties.keywords = "usinas, energia, criptoagilidade, QKD, proxy, NIST"

doc.save(OUTPUT)
print(OUTPUT)
